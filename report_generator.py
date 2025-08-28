#report_generator.py
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import io
import re
import os

# python-docx가 있으면 import, 없으면 None
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    WD_ALIGN_PARAGRAPH = None

def register_korean_font():
    """한국어 폰트 등록 - 여러 옵션 시도"""
    font_options = [
        'NOTOSANSKR-VF.TTF',
        'NanumGothicCoding.ttf',
        'NanumGothicCoding-Bold.ttf',
        'malgun.ttf',  # Windows 기본 폰트
        'gulim.ttc',   # Windows 기본 폰트
    ]
    
    for font_file in font_options:
        try:
            if os.path.exists(font_file):
                pdfmetrics.registerFont(TTFont('KoreanFont', font_file))
                print(f"한국어 폰트 등록 성공: {font_file}")
                return True
        except Exception as e:
            print(f"폰트 등록 실패 ({font_file}): {e}")
            continue
    
    # 폰트 파일이 없으면 기본 폰트 사용
    print("한국어 폰트를 찾을 수 없어 기본 폰트를 사용합니다.")
    return False

def clean_text_for_pdf(text):
    """PDF용 텍스트 정리 - HTML 태그 제거 및 안전한 형식으로 변환 - 개선된 버전"""
    if not text:
        return ""
    
    # HTML 태그 제거
    text = re.sub(r'<br\s*/?>', '\n', text)  # <br> 태그를 줄바꿈으로
    text = re.sub(r'<[^>]+>', '', text)  # 모든 HTML 태그 제거
    
    # 특수 문자 처리
    text = text.replace('•', '•')  # bullet point
    text = text.replace('–', '-')  # en dash
    text = text.replace('—', '-')  # em dash
    text = text.replace('"', '"')  # smart quotes
    text = text.replace('"', '"')  # smart quotes
    text = text.replace(''', "'")  # smart apostrophe
    text = text.replace(''', "'")  # smart apostrophe
    
    # 표 관련 특수 문자 처리
    text = text.replace('│', '|')  # box drawing characters
    text = text.replace('┌', '')
    text = text.replace('┐', '')
    text = text.replace('└', '')
    text = text.replace('┘', '')
    text = text.replace('├', '')
    text = text.replace('┤', '')
    text = text.replace('┬', '')
    text = text.replace('┴', '')
    text = text.replace('─', '-')
    
    # 연속된 공백 정리 (표 셀 내에서는 보존)
    if '|' not in text:  # 표가 아닌 경우에만 공백 정리
        text = re.sub(r'\s+', ' ', text)
    
    # 줄바꿈 정리 (표가 아닌 경우에만)
    if '|' not in text:
        text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

def parse_table_from_text(text):
    """텍스트에서 표 형식을 파싱하여 2D 배열로 변환 - 오류 처리 강화"""
    try:
        if not text or not isinstance(text, str):
            return [], None
            
        lines = text.strip().split('\n')
        table_data = []
        table_title = None
        
        # 표 제목 찾기 (표 위의 텍스트)
        title_lines = []
        table_started = False
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # 표 시작 확인
            if is_table_row(line):
                table_started = True
                # 이전까지의 텍스트를 제목으로 처리
                if title_lines:
                    table_title = ' '.join(title_lines).strip()
                    # 제목에서 불필요한 문자 제거
                    table_title = re.sub(r'^\*\*|\*\*$', '', table_title)  # 마크다운 볼드 제거
                    table_title = re.sub(r'^#+\s*', '', table_title)  # 마크다운 헤더 제거
                break
            else:
                # 표가 시작되기 전까지의 텍스트를 제목으로 저장
                # 단, 너무 긴 텍스트는 제목이 아닐 수 있음
                if len(line) < 100:  # 100자 이하만 제목으로 간주
                    title_lines.append(line)
        
        # 표 데이터 파싱
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 구분선 제거 (마크다운 표 구분선)
            if re.match(r'^[\s\-=_:|]+\s*$', line):
                continue
                
            # 표 행인지 확인
            if is_table_row(line):
                try:
                    cells = parse_table_row(line)
                    if cells:
                        table_data.append(cells)
                except Exception as e:
                    # 개별 행 파싱 오류 시 건너뛰기
                    print(f"행 파싱 오류: {e}")
                    continue
        
        # 표 데이터 정규화 (모든 행이 같은 열 수를 가지도록)
        if table_data:
            try:
                max_cols = max(len(row) for row in table_data)
                normalized_data = []
                for row in table_data:
                    # 부족한 열은 빈 문자열로 채움
                    normalized_row = row + [''] * (max_cols - len(row))
                    normalized_data.append(normalized_row)
                return normalized_data, table_title
            except Exception as e:
                print(f"표 데이터 정규화 오류: {e}")
                return table_data, table_title
        
        return table_data, table_title
        
    except Exception as e:
        print(f"표 파싱 중 오류 발생: {e}")
        return [], None

def is_table_row(line):
    """한 줄이 표 행인지 확인"""
    # | 구분자가 있는 경우
    if '|' in line:
        return True
    
    # 탭으로 구분된 경우
    if '\t' in line:
        return True
    
    # 2개 이상의 공백으로 구분된 경우 (정렬된 텍스트)
    if re.search(r'\s{2,}', line):
        return True
    
    return False

def parse_table_row(line):
    """표 행을 파싱하여 셀 배열로 변환 - 오류 처리 강화"""
    try:
        if not line or not isinstance(line, str):
            return []
        
        # | 구분자로 분할 (마크다운 표 형식)
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            # 첫 번째와 마지막 빈 셀 제거 (마크다운 표 형식)
            if cells and not cells[0].strip():
                cells = cells[1:]
            if cells and not cells[-1].strip():
                cells = cells[:-1]
            return cells
        
        # 탭으로 구분된 경우
        elif '\t' in line:
            cells = [cell.strip() for cell in line.split('\t') if cell.strip()]
            return cells
        
        # 2개 이상의 공백으로 구분된 경우
        elif re.search(r'\s{2,}', line):
            cells = [cell.strip() for cell in re.split(r'\s{2,}', line) if cell.strip()]
            return cells
        
        return []
        
    except Exception as e:
        print(f"행 파싱 오류: {e}")
        return []

def is_table_format(text):
    """텍스트가 표 형식인지 확인 - 오류 처리 강화"""
    try:
        if not text or not isinstance(text, str):
            return False
            
        lines = text.strip().split('\n')
        if len(lines) < 2:
            return False
        
        # 표 구분자 확인
        table_indicators = ['|', '\t']
        table_line_count = 0
        
        for line in lines[:5]:  # 처음 5줄만 확인
            if any(indicator in line for indicator in table_indicators):
                table_line_count += 1
        
        # 2줄 이상에 표 구분자가 있으면 표로 인식
        if table_line_count >= 2:
            return True
        
        # 구분선 확인 (마크다운 표 구분선)
        for line in lines:
            if re.match(r'^[\s\-=_:|]+\s*$', line.strip()):
                return True
        
        # 정렬된 텍스트 확인 (2개 이상의 공백으로 구분)
        aligned_line_count = 0
        for line in lines[:3]:
            if re.search(r'\s{2,}', line.strip()):
                aligned_line_count += 1
        
        if aligned_line_count >= 2:
            return True
        
        return False
        
    except Exception as e:
        print(f"표 형식 확인 오류: {e}")
        return False

def create_table_with_improved_style(table_data, font_registered):
    """개선된 스타일로 표 생성 - 내용에 맞는 자동 크기 조절"""
    if not table_data or len(table_data) == 0:
        return None
    
    try:
        # 데이터 검증
        if not isinstance(table_data, list) or len(table_data) == 0:
            return None
        
        # 모든 행이 리스트인지 확인
        for row in table_data:
            if not isinstance(row, list):
                return None
        
        # 헤더 행 확인
        has_header = is_header_row(table_data[0]) if table_data else False
        
        # 열 수 확인 및 정규화
        max_cols = max(len(row) for row in table_data) if table_data else 1
        if max_cols == 0:
            return None
        
        # 데이터 정규화 (모든 행이 같은 열 수를 가지도록)
        normalized_data = []
        for row in table_data:
            normalized_row = row + [''] * (max_cols - len(row))
            normalized_data.append(normalized_row)
        
        # 페이지 너비 설정
        page_width = 8.27 * inch  # A4 너비
        margin = 0.5 * inch
        available_width = page_width - 2 * margin
        
        # 각 열의 내용 길이를 분석하여 적절한 너비 계산
        col_widths = []
        min_col_width = available_width / max_cols * 0.3  # 최소 너비 (30%)
        max_col_width = available_width / max_cols * 2.0  # 최대 너비 (200%)
        
        for col_idx in range(max_cols):
            max_content_length = 0
            for row in normalized_data:
                if col_idx < len(row):
                    content = str(row[col_idx]) if row[col_idx] is not None else ""
                    # 대략적인 문자 길이 계산 (한글은 2배로 계산)
                    char_count = len(content)
                    korean_count = len([c for c in content if ord(c) > 127])
                    adjusted_length = char_count + korean_count
                    max_content_length = max(max_content_length, adjusted_length)
            
            # 내용 길이에 따른 너비 계산
            if max_content_length > 0:
                # 기본 폰트 크기 8pt 기준으로 계산 (더 정확한 계산)
                # 한글 문자는 약 1.2배, 영문/숫자는 약 0.6배 너비
                estimated_width = max_content_length * 8 * 0.7  # 조정된 계수
                col_width = max(min_col_width, min(estimated_width, max_col_width))
            else:
                col_width = min_col_width
            
            col_widths.append(col_width)
        
        # 최소 너비 보장
        for i in range(len(col_widths)):
            if col_widths[i] < min_col_width:
                col_widths[i] = min_col_width
        
        # 전체 너비가 페이지를 넘지 않도록 조정
        total_width = sum(col_widths)
        if total_width > available_width:
            scale_factor = available_width / total_width
            col_widths = [w * scale_factor for w in col_widths]
        
        # 각 셀을 Paragraph 객체로 변환 (자동 줄바꿈 지원)
        from reportlab.lib.styles import getSampleStyleSheet
        styles = getSampleStyleSheet()
        
        # 셀 스타일 생성 (자동 줄바꿈 강화)
        if font_registered:
            cell_style = ParagraphStyle(
                'TableCell',
                parent=styles['Normal'],
                fontName='KoreanFont',
                fontSize=8,
                alignment=0,  # 0=LEFT, 1=CENTER, 2=RIGHT
                leftIndent=0,
                rightIndent=0,
                wordWrap='CJK',  # 한글 자동 줄바꿈
                spaceBefore=2,
                spaceAfter=2,
                leading=10,  # 줄 간격
                splitLongWords=True,  # 긴 단어 분할
                hyphenationLang='ko_KR'  # 한글 하이픈
            )
        else:
            cell_style = ParagraphStyle(
                'TableCell',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=8,
                alignment=0,  # 0=LEFT, 1=CENTER, 2=RIGHT
                leftIndent=0,
                rightIndent=0,
                wordWrap='CJK',  # 한글 자동 줄바꿈
                spaceBefore=2,
                spaceAfter=2,
                leading=10,  # 줄 간격
                splitLongWords=True,  # 긴 단어 분할
            )
        
        # 표 데이터를 Paragraph 객체로 변환
        paragraph_data = []
        for row in normalized_data:
            paragraph_row = []
            for cell in row:
                try:
                    # 셀 내용을 안전하게 문자열로 변환
                    cell_text = str(cell) if cell is not None else ""
                    # HTML 태그 제거
                    cell_text = clean_text_for_pdf(cell_text)
                    # 셀 내용을 Paragraph로 변환
                    paragraph_cell = Paragraph(cell_text, cell_style)
                    paragraph_row.append(paragraph_cell)
                except Exception as e:
                    # 개별 셀 오류 시 빈 셀로 대체
                    paragraph_cell = Paragraph("", cell_style)
                    paragraph_row.append(paragraph_cell)
            paragraph_data.append(paragraph_row)
        
        # 표 생성 (Paragraph 객체 사용, 자동 너비 조절)
        table = Table(paragraph_data, colWidths=col_widths)
        
        # 기본 스타일 (격자만)
        table_style = TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('WORDWRAP', (0, 0), (-1, -1), True),  # 자동 줄바꿈 활성화
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#F8F9FA')]),  # 줄무늬 배경
        ])
        
        # 헤더 스타일 추가
        if has_header and len(paragraph_data) > 0:
            try:
                table_style.add('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86AB'))
                table_style.add('TEXTCOLOR', (0, 0), (-1, 0), colors.white)
                table_style.add('FONTWEIGHT', (0, 0), (-1, 0), 'bold')
            except Exception as e:
                # 헤더 스타일링 실패 시 기본 스타일만 적용
                pass
        
        table.setStyle(table_style)
        return table
        
    except Exception as e:
        print(f"표 생성 중 오류 발생: {e}")
        # 오류 발생 시 간단한 텍스트 표로 대체
        try:
            simple_data = []
            for row in table_data:
                simple_row = [str(cell) if cell is not None else "" for cell in row]
                simple_data.append(simple_row)
            
            if simple_data:
                # 간단한 표도 자동 크기 조절 적용
                max_cols = max(len(row) for row in simple_data) if simple_data else 1
                page_width = 8.27 * inch
                margin = 0.5 * inch
                available_width = page_width - 2 * margin
                
                # 간단한 너비 계산
                col_width = available_width / max_cols
                col_widths = [col_width] * max_cols
                
                simple_table = Table(simple_data, colWidths=col_widths)
                simple_table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('WORDWRAP', (0, 0), (-1, -1), True),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ]))
                return simple_table
        except:
            pass
        
        return None

def is_header_row(row):
    """행이 헤더인지 확인"""
    if not row:
        return False
    
    # 헤더로 보이는 키워드들
    header_keywords = [
        '항목', '구분', '분류', '종류', '유형', '타입', '카테고리',
        '특성', '특징', '속성', '성질', '성격',
        '근거', '이유', '원인', '배경', '기반',
        '내용', '설명', '상세', '세부',
        '비율', '퍼센트', '수치', '값', '데이터',
        '날짜', '기간', '시기', '연도', '월', '일',
        '이름', '명칭', '제목', '표제',
        '번호', '순서', '순번', '인덱스'
    ]
    
    # 행의 모든 셀을 확인
    for cell in row:
        cell_lower = cell.lower().strip()
        for keyword in header_keywords:
            if keyword in cell_lower:
                return True
    
    # 셀 내용이 짧고 명확한 경우 (헤더일 가능성)
    short_cells = sum(1 for cell in row if len(cell.strip()) <= 10)
    if short_cells >= len(row) * 0.7:  # 70% 이상이 짧은 경우
        return True
    
    return False

def generate_pdf_report(content, user_inputs):
    """PDF 보고서 생성 - 표 처리 개선 및 폰트 문제 해결"""
    
    # 메모리 버퍼에 PDF 생성
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    
    # 스타일 설정
    styles = getSampleStyleSheet()
    
    # 한국어 폰트 등록
    font_registered = register_korean_font()
    
    # 커스텀 스타일 생성
    if font_registered:
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName='KoreanFont',
            fontSize=16,
            spaceAfter=12,
            alignment=1  # 중앙 정렬
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName='KoreanFont',
            fontSize=14,
            spaceAfter=8
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='KoreanFont',
            fontSize=10,
            spaceAfter=6
        )
    else:
        # 기본 폰트 사용 (한국어 지원 안됨)
        title_style = styles['Heading1']
        heading_style = styles['Heading2']
        normal_style = styles['Normal']
    
    # 내용을 문단으로 분할
    story = []
    
    # 제목 추가
    project_name = user_inputs.get('project_name', '프로젝트')
    title_text = f"{project_name} 분석 보고서"
    story.append(Paragraph(title_text, title_style))
    story.append(Spacer(1, 20))
    
    # 내용 파싱 및 추가
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # 표 형식 처리
        if is_table_format(para):
            try:
                table_data, table_title = parse_table_from_text(para)
                if table_data and len(table_data) > 0:
                    # 표 제목이 있으면 먼저 추가 (HTML 태그 제거)
                    if table_title:
                        story.append(Paragraph(clean_text_for_pdf(table_title), heading_style))
                        story.append(Spacer(1, 6))
                    
                    # 표 생성 - 표 데이터는 clean_text_for_pdf를 거치지 않음!
                    table = create_table_with_improved_style(table_data, font_registered)
                    
                    if table:
                        story.append(table)
                        story.append(Spacer(1, 12))
                    else:
                        # 표 생성 실패 시 일반 텍스트로 변환
                        if table_title:
                            story.append(Paragraph(f"[표 제목: {clean_text_for_pdf(table_title)}]", normal_style))
                        story.append(Paragraph(f"[표 데이터: {clean_text_for_pdf(para[:200])}...]", normal_style))
                        story.append(Spacer(1, 6))
                else:
                    # 표 데이터가 없으면 일반 텍스트로 처리
                    story.append(Paragraph(clean_text_for_pdf(para), normal_style))
                    story.append(Spacer(1, 6))
            except Exception as e:
                # 표 처리 중 오류 발생 시 일반 텍스트로 변환
                print(f"표 처리 오류: {e}")
                story.append(Paragraph(clean_text_for_pdf(para), normal_style))
                story.append(Spacer(1, 6))
            continue
        
        # 일반 텍스트 처리
        lines = para.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 제목 처리
            if line.startswith('# '):
                text = clean_text_for_pdf(line[2:].strip())
                story.append(Paragraph(text, title_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                text = clean_text_for_pdf(line[3:].strip())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 8))
            elif line.startswith('### '):
                text = clean_text_for_pdf(line[4:].strip())
                story.append(Paragraph(text, heading_style))
                story.append(Spacer(1, 6))
            elif line.startswith('---'):
                story.append(Spacer(1, 12))
            else:
                # 일반 텍스트 처리
                if line:
                    clean_line = clean_text_for_pdf(line)
                    if clean_line:
                        story.append(Paragraph(clean_line, normal_style))
    
    # PDF 생성
    try:
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"PDF 생성 오류: {e}")
        # 오류 발생 시 간단한 텍스트로 재시도
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        simple_story = []
        
        # 간단한 텍스트로 변환
        simple_story.append(Paragraph(f"{project_name} 분석 보고서", title_style))
        simple_story.append(Spacer(1, 20))
        
        # 원본 텍스트를 그대로 사용 (HTML 태그 제거)
        clean_content = clean_text_for_pdf(content)
        paragraphs = clean_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                simple_story.append(Paragraph(para.strip(), normal_style))
                simple_story.append(Spacer(1, 6))
        
        doc.build(simple_story)
        buffer.seek(0)
        return buffer.getvalue()

def generate_word_report(content, user_inputs):
    """Word 문서 보고서 생성 - 표 처리 개선"""
    
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 모듈이 설치되지 않았습니다. 'pip install python-docx'로 설치해주세요.")
    
    # Word 문서 생성
    doc = Document()
    
    # 제목 설정
    project_name = user_inputs.get('project_name', '프로젝트')
    title = doc.add_heading(f"{project_name} 분석 보고서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 내용 파싱 및 추가
    paragraphs = content.split('\n\n')
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # 표 형식 처리
        if is_table_format(para):
            table_data, table_title = parse_table_from_text(para)
            if table_data and len(table_data) > 0:
                try:
                    # 표 제목이 있으면 먼저 추가
                    if table_title:
                        doc.add_heading(clean_text_for_pdf(table_title), level=3)
                    
                    # 헤더 행 확인
                    has_header = is_header_row(table_data[0]) if table_data else False
                    
                    # Word 표 생성
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    
                    # 표 스타일 개선 - 자동 크기 조절
                    table.allow_autofit = True
                    table.autofit = True
                    
                    # 데이터 채우기
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            if i < len(table.rows) and j < len(table.rows[i].cells):
                                cell_text = clean_text_for_pdf(cell)
                                table.rows[i].cells[j].text = cell_text
                                
                                # 셀 자동 줄바꿈 설정
                                cell_obj = table.rows[i].cells[j]
                                for paragraph in cell_obj.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    # 자동 줄바꿈을 위한 설정
                                    paragraph.paragraph_format.space_before = 0
                                    paragraph.paragraph_format.space_after = 0
                                
                                # 헤더 행 스타일링
                                if has_header and i == 0:
                                    cell_obj = table.rows[i].cells[j]
                                    for paragraph in cell_obj.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                                            run.font.color.rgb = None  # 기본 색상
                
                except Exception as e:
                    print(f"표 생성 오류: {e}")
                    # 표 생성 실패 시 일반 텍스트로 변환
                    if table_title:
                        doc.add_paragraph(f"[표 제목: {table_title}]")
                    doc.add_paragraph(f"[표 데이터: {para[:100]}...]")
                
                doc.add_paragraph()  # 표 후 빈 줄
                continue
        
        # 일반 텍스트 처리
        lines = para.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 제목 처리
            if line.startswith('# '):
                text = clean_text_for_pdf(line[2:].strip())
                doc.add_heading(text, level=1)
            elif line.startswith('## '):
                text = clean_text_for_pdf(line[3:].strip())
                doc.add_heading(text, level=2)
            elif line.startswith('### '):
                text = clean_text_for_pdf(line[4:].strip())
                doc.add_heading(text, level=3)
            elif line.startswith('---'):
                doc.add_paragraph()  # 빈 줄 추가
            else:
                # 일반 텍스트 처리
                if line:
                    clean_line = clean_text_for_pdf(line)
                    if clean_line:
                        doc.add_paragraph(clean_line)
    
    # 메모리 버퍼에 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_report_content(report_type, include_charts, include_recommendations, include_appendix):
    """보고서 내용 생성 - 보고서 유형별 차이점 적용"""
    from user_state import get_user_inputs
    user_inputs = get_user_inputs()
    
    # 기본 정보
    report_content = f"""
# {user_inputs.get('project_name', '프로젝트')} 분석 보고서
**보고서 유형**: {report_type}

## 📋 프로젝트 기본 정보
- **프로젝트명**: {user_inputs.get('project_name', 'N/A')}
- **건축주**: {user_inputs.get('owner', 'N/A')}
- **대지위치**: {user_inputs.get('site_location', 'N/A')}
- **대지면적**: {user_inputs.get('site_area', 'N/A')}
- **건물용도**: {user_inputs.get('building_type', 'N/A')}
- **프로젝트 목표**: {user_inputs.get('project_goal', 'N/A')}

"""
    
    # 분석 결과 추가
    import streamlit as st
    if st.session_state.get('cot_history'):
        if report_type == "전체 분석 보고서":
            # 전체 분석 보고서: 모든 상세 내용 포함
            report_content += "## 전체 분석 결과\n"
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

{history.get('result', '')}

---
"""
        
        elif report_type == "요약 보고서":
            # 요약 보고서: 핵심 요약과 인사이트만
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**핵심 요약**: {history.get('summary', '')}

**주요 인사이트**: {history.get('insight', '')}

---
"""
        
        elif report_type == "전문가 보고서":
            # 전문가 보고서: 기술적 분석과 전문적 권장사항
            report_content += "## 전문가 분석 결과\n"
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**분석 요약**: {history.get('summary', '')}

**전문가 인사이트**: {history.get('insight', '')}

**기술적 분석**:
{history.get('result', '')[:500]}...

---
"""
        
        elif report_type == "클라이언트 보고서":
            # 클라이언트 보고서: 비즈니스 관점의 핵심 내용
            report_content += "## 💼 비즈니스 분석 결과\n"
            for i, history in enumerate(st.session_state.cot_history, 1):
                report_content += f"""
### {i}. {history['step']}

**비즈니스 요약**: {history.get('summary', '')}

**핵심 가치**: {history.get('insight', '')}

**실행 가능한 제안**:
{history.get('result', '')[:300]}...

---
"""
    
    # 추가 섹션들 (보고서 유형별로 다르게 적용)
    if include_charts:
        if report_type == "전체 분석 보고서":
            report_content += """
## 상세 차트 및 다이어그램
(모든 차트 및 다이어그램이 포함됩니다)
"""
        elif report_type == "요약 보고서":
            report_content += """
## 핵심 차트
(주요 차트만 포함됩니다)
"""
        elif report_type == "전문가 보고서":
            report_content += """
## 전문가 차트 및 분석 다이어그램
(기술적 분석을 위한 상세 차트가 포함됩니다)
"""
        elif report_type == "클라이언트 보고서":
            report_content += """
## 💼 비즈니스 차트
(비즈니스 의사결정을 위한 핵심 차트가 포함됩니다)
"""
    
    if include_recommendations:
        if report_type == "전체 분석 보고서":
            report_content += """
## 💡 종합 권장사항
분석 결과를 바탕으로 한 상세한 권장사항이 포함됩니다.
"""
        elif report_type == "요약 보고서":
            report_content += """
## 💡 핵심 권장사항
가장 중요한 권장사항만 포함됩니다.
"""
        elif report_type == "전문가 보고서":
            report_content += """
## 💡 전문가 권장사항
기술적 관점에서의 전문적 권장사항이 포함됩니다.
"""
        elif report_type == "클라이언트 보고서":
            report_content += """
## 💡 비즈니스 권장사항
비즈니스 관점에서의 실행 가능한 권장사항이 포함됩니다.
"""
    
    if include_appendix:
        if report_type == "전체 분석 보고서":
            report_content += """
## 📋 상세 부록
모든 추가 자료 및 참고문헌이 포함됩니다.
"""
        elif report_type == "요약 보고서":
            report_content += """
## 📋 핵심 부록
주요 참고자료만 포함됩니다.
"""
        elif report_type == "전문가 보고서":
            report_content += """
## 📋 전문가 부록
기술적 참고자료 및 전문 문헌이 포함됩니다.
"""
        elif report_type == "클라이언트 보고서":
            report_content += """
## 📋 비즈니스 부록
비즈니스 관련 참고자료가 포함됩니다.
"""
    
    return report_content