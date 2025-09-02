"""
분석 시스템 핵심 구조 UI
- 용도/목적 분류
- 단계 자동 제안
- 필수 단계 포함
- 번외 항목 추가
- 순서 변경 및 추가/삭제
- 전체 순서 확정 및 분석 실행
"""

import streamlit as st
import json
import time
from datetime import datetime
from user_state import get_user_inputs, save_step_result, append_step_history
from analysis_system import (
    AnalysisSystem, PurposeType, ObjectiveType, AnalysisWorkflow
)
from agent_executor import (
    run_requirement_table,
    run_ai_reasoning,
    run_precedent_comparison,
    run_strategy_recommendation,
    run_requirement_analyzer,
    run_document_analyzer,
    # 하이데라바드 프로젝트 전용 함수들 (일시적으로 비활성화)
    # run_hyderabad_campus_expansion_analysis,
    # run_hyderabad_research_infra_strategy,
    # run_hyderabad_talent_collaboration_infra,
    # run_hyderabad_welfare_branding_environment,
    # run_hyderabad_security_zoning_plan,
    # run_hyderabad_masterplan_roadmap,
)
from utils_pdf import (
    initialize_vector_system,
    extract_text_from_pdf,
    get_pdf_summary,
)
from dsl_to_prompt import convert_dsl_to_prompt

# 파일 상단에 상수 정의
REQUIRED_FIELDS = ["project_name", "building_type", "site_location", "owner", "site_area", "project_goal"]
FEEDBACK_TYPES = ["추가 분석 요청", "수정 요청", "다른 관점 제시", "구조 변경", "기타"]

def execute_claude_analysis(prompt, description):
    """Claude 분석 실행 함수 - 세션 상태 기반 모델 선택"""
    
    # 세션 상태에서 선택된 모델 가져오기
    selected_model = st.session_state.get('selected_model', 'claude-sonnet-4-20250514')
    
    # SDK 방식으로 실행 (DSPy 설정 변경 없이) - 재시도 로직 포함
    from init_dspy import execute_with_sdk_with_retry
    
    # 진행 상황 표시
    with st.spinner(f"{description} 분석 중... (재시도 로직 포함)"):
        result = execute_with_sdk_with_retry(prompt, selected_model, max_retries=3)
    
    # 오류 메시지 개선
    if result.startswith("❌") or result.startswith("⚠️"):
        st.error(f"분석 중 오류가 발생했습니다: {result}")
        # 사용자에게 재시도 옵션 제공
        if st.button("다시 시도", key=f"retry_{description}"):
            st.rerun()
    
    return result

def create_analysis_workflow(purpose_enum, objective_enums):
    """워크플로우 생성 함수"""
    system = AnalysisSystem()
    return system.suggest_analysis_steps(purpose_enum, objective_enums)

def validate_user_inputs(user_inputs):
    """사용자 입력 검증 함수"""
    missing_fields = [field for field in REQUIRED_FIELDS if not user_inputs.get(field)]
    return missing_fields

def create_pdf_summary_dict(user_inputs, pdf_summary):
    """PDF 요약 딕셔너리 생성 함수"""
    return {
        "pdf_summary": pdf_summary,
        "project_name": user_inputs.get("project_name", ""),
        "owner": user_inputs.get("owner", ""),
        "site_location": user_inputs.get("site_location", ""),
        "site_area": user_inputs.get("site_area", ""),
        "building_type": user_inputs.get("building_type", ""),
        "project_goal": user_inputs.get("project_goal", "")
    }

def render_purpose_selection():
    """1단계: 용도 선택"""
    st.subheader("1단계: 건물 용도 선택")
    
    purpose_options = [purpose.value for purpose in PurposeType]
    selected_purpose = st.selectbox(
        "건물 용도를 선택하세요",
        purpose_options,
        key="selected_purpose"
    )
    
    if selected_purpose:
        return PurposeType(selected_purpose)
    return None

def render_objective_selection(purpose: PurposeType, system: AnalysisSystem):
    """2단계: 목적 선택"""
    st.subheader("2단계: 분석 목적 선택")
    
    available_objectives = system.get_available_objectives(purpose)
    objective_options = [obj.value for obj in available_objectives]
    
    selected_objectives = st.multiselect(
        "분석 목적을 선택하세요 (복수 선택 가능)",
        objective_options,
        key="selected_objectives"
    )
    
    if selected_objectives:
        return [ObjectiveType(obj) for obj in selected_objectives]
    return []

def render_analysis_steps_management(selected_purpose, selected_objectives, system):
    """3단계: 분석 단계 관리"""
    st.subheader("3단계: 분석 단계 관리")
    
    # 워크플로우 생성 - suggest_analysis_steps 사용
    workflow = system.suggest_analysis_steps(selected_purpose, selected_objectives)
    
    # 제거된 단계들을 필터링
    removed_steps = st.session_state.get('removed_steps', set())
    workflow.steps = [step for step in workflow.steps if step.id not in removed_steps]
    workflow.custom_steps = [step for step in workflow.custom_steps if step.id not in removed_steps]
    
    # 최종 워크플로우 가져오기
    final_steps = system.get_final_workflow(workflow)
    
    st.markdown("### 현재 분석 단계들:")
    
    # 모든 단계를 하나의 리스트로 통합
    all_steps = final_steps.copy()
    
    for i, step in enumerate(all_steps):
        col1, col2, col3, col4, col5 = st.columns([3, 1, 1, 1, 1])
        
        with col1:
            if step.is_required:
                st.markdown(f" **{step.title}** (필수)")
            elif step.is_recommended:
                st.markdown(f" **{step.title}** (권장)")
            else:
                st.markdown(f" **{step.title}** (선택)")
            st.markdown(f"*{step.description}*")
        
        with col2:
            # 제거 버튼
            if not step.is_required:
                if st.button("제거", key=f"remove_{step.id}_{i}", use_container_width=True):
                    # 제거된 단계 세트에 추가
                    if 'removed_steps' not in st.session_state:
                        st.session_state.removed_steps = set()
                    st.session_state.removed_steps.add(step.id)
                    st.success(f"'{step.title}' 단계가 제거되었습니다!")
                    st.rerun()
            else:
                st.markdown("제거")
        
        with col3:
            # 위로 이동 버튼
            if i > 0:
                if st.button("위로", key=f"up_{step.id}_{i}", use_container_width=True):
                    # 현재 단계와 위 단계의 순서를 바꿈
                    all_steps[i], all_steps[i-1] = all_steps[i-1], all_steps[i]
                    
                    # 순서 번호 업데이트
                    for j, s in enumerate(all_steps):
                        s.order = (j + 1) * 10
                    
                    # 워크플로우 업데이트
                    workflow.steps = [s for s in all_steps if s.is_required or s.is_recommended]
                    workflow.custom_steps = [s for s in all_steps if s.is_optional]
                    
                    # 세션 상태에 저장
                    st.session_state.workflow_steps = all_steps
                    st.session_state.current_workflow = workflow
                    
                    st.success(f"'{step.title}' 단계가 위로 이동되었습니다!")
                    st.rerun()
            else:
                st.markdown("위로")
        
        with col4:
            # 아래로 이동 버튼
            if i < len(all_steps) - 1:
                if st.button("아래로", key=f"down_{step.id}_{i}", use_container_width=True):
                    # 현재 단계와 아래 단계의 순서를 바꿈
                    all_steps[i], all_steps[i+1] = all_steps[i+1], all_steps[i]
                    
                    # 순서 번호 업데이트
                    for j, s in enumerate(all_steps):
                        s.order = (j + 1) * 10
                    
                    # 워크플로우 업데이트
                    workflow.steps = [s for s in all_steps if s.is_required or s.is_recommended]
                    workflow.custom_steps = [s for s in all_steps if s.is_optional]
                    
                    # 세션 상태에 저장
                    st.session_state.workflow_steps = all_steps
                    st.session_state.current_workflow = workflow
                    
                    st.success(f"'{step.title}' 단계가 아래로 이동되었습니다!")
                    st.rerun()
            else:
                st.markdown("아래로")
        
        with col5:
            st.markdown(f"**{i+1}**")
    
    # 순서 재정렬 버튼
    if st.button("전체 순서 재정렬", key="reorder_all", use_container_width=True):
        # 모든 단계를 10단위로 재정렬
        for i, step in enumerate(all_steps):
            step.order = (i + 1) * 10
        
        # 워크플로우 업데이트
        workflow.steps = [s for s in all_steps if s.is_required or s.is_recommended]
        workflow.custom_steps = [s for s in all_steps if s.is_optional]
        
        # 세션 상태에 저장
        st.session_state.workflow_steps = all_steps
        st.session_state.current_workflow = workflow
        
        st.success("순서가 재정렬되었습니다!")
        st.rerun()
    
    # 워크플로우를 세션 상태에 저장
    st.session_state.workflow_steps = all_steps
    st.session_state.current_workflow = workflow
    
    return workflow

def render_workflow_summary(workflow: AnalysisWorkflow, system: AnalysisSystem):
    """4단계: 워크플로우 요약"""
    st.subheader("4단계: 최종 분석 워크플로우")
    
    st.markdown(f"**선택된 용도:** {workflow.purpose.value}")
    st.markdown(f"**선택된 목적:** {workflow.objective.value}")
    
    st.markdown("### 최종 분석 단계들:")
    
    final_steps = system.get_final_workflow(workflow)
    
    # 각 단계별 웹 검색 설정을 저장할 딕셔너리 초기화
    if 'web_search_settings' not in st.session_state:
        st.session_state.web_search_settings = {}
    
    for i, step in enumerate(final_steps, 1):
        if step.is_required:
            level_icon = "��"
            level_text = "필수"
        elif step.is_recommended:
            level_icon = "��"
            level_text = "권장"
        else:
            level_icon = "��"
            level_text = "선택"
        
        # 각 단계별 웹 검색 체크박스
        col1, col2 = st.columns([4, 1])
        
        with col1:
            st.markdown(f"{i}. {level_icon} **{step.title}** ({level_text})")
            st.markdown(f"   - {step.description}")
        
        with col2:
            # 웹 검색 체크박스 (기본값: False)
            web_search_key = f"web_search_{step.id}"
            if web_search_key not in st.session_state.web_search_settings:
                st.session_state.web_search_settings[web_search_key] = False
            
            st.session_state.web_search_settings[web_search_key] = st.checkbox(
                "웹 검색",
                value=st.session_state.web_search_settings[web_search_key],
                key=web_search_key
            )
    
    # 분석 실행 버튼을 여기서 직접 처리
    if st.button("분석 실행", type="primary", use_container_width=True, key="execute_analysis"):
        # 분석 상태 설정
        st.session_state.analysis_started = True
        # current_step_index를 0으로 초기화하지 않고 기존 값 유지
        if 'current_step_index' not in st.session_state:
            st.session_state.current_step_index = 0
        # cot_history를 초기화하지 않고 기존 값 유지
        if 'cot_history' not in st.session_state:
            st.session_state.cot_history = []
        st.session_state.workflow_steps = final_steps
        st.session_state.show_feedback = False
        
        st.success("✅ 분석이 시작되었습니다! 각 단계를 수동으로 진행하세요.")
        st.rerun()
    
    return False  # 버튼 반환 대신 False 반환

def validate_prompt_structure(dsl_block: dict) -> bool:
    """DSL 블록의 구조 유효성 검증"""
    content_dsl = dsl_block.get("content_dsl", {})
    output_structure = content_dsl.get("output_structure", [])
    
    if not output_structure:
        st.warning("⚠️ 이 블록에는 output_structure가 정의되지 않았습니다.")
        return False
    
    if len(output_structure) < 1:
        st.warning("⚠️ output_structure가 비어있습니다.")
        return False
    
    return True

def debug_analysis_result(result: str, output_structure: list):
    """분석 결과 디버깅 정보 표시"""
    with st.expander("디버깅 정보", expanded=False):
        st.markdown("**Output Structure:**")
        for i, structure in enumerate(output_structure, 1):
            st.markdown(f"{i}. {structure}")
        
        st.markdown("**AI 응답 길이:**")
        st.markdown(f"{len(result)} 문자")
        
        st.markdown("**구조명 매칭 결과:**")
        for structure in output_structure:
            found = structure in result
            st.markdown(f"- {structure}: {'✅' if found else '❌'}")
        
        st.markdown("**전체 응답 미리보기:**")
        st.code(result[:500] + "..." if len(result) > 500 else result)

def render_analysis_execution():
    """분석 실행 UI - 단계별 진행 방식"""
    if not st.session_state.get('analysis_started', False):
        return

    st.title("건축 분석 워크플로우")
    st.subheader("분석 실행")

    # cot_history 디버깅 정보 추가
    st.sidebar.markdown("분석 실행 디버깅")
    if st.session_state.get('cot_history'):
        st.sidebar.write(f"**완료된 분석: {len(st.session_state.cot_history)}개**")
        for i, history in enumerate(st.session_state.cot_history):
            step_name = history.get('step', f'단계 {i+1}')
            result_length = len(history.get('result', ''))
            st.sidebar.write(f"**{i+1}. {step_name}**")
            st.sidebar.write(f"   길이: {result_length} 문자")
            if result_length > 0:
                preview = history.get('result', '')[:30] + "..." if result_length > 30 else history.get('result', '')
                st.sidebar.write(f"   미리보기: {preview}")
    else:
        st.sidebar.write("**완료된 분석: 없음**")

    # 디버깅 정보 추가 (render_analysis_execution 함수 시작 부분에)
    st.sidebar.markdown("상태 디버깅")
    st.sidebar.write(f"**현재 단계**: {st.session_state.get('current_step_index', 0) + 1}")
    st.sidebar.write(f"**분석 시작됨**: {st.session_state.get('analysis_started', False)}")
    st.sidebar.write(f"**완료된 분석**: {len(st.session_state.get('cot_history', []))}개")
    st.sidebar.write(f"**현재 단계 제목**: {st.session_state.get('current_step', {}).get('title', '')}")
    if st.session_state.get('current_step'):
        st.sidebar.write(f"**현재 블록 제목**: {st.session_state.get('current_step').get('title', '')}")
        st.sidebar.write(f"**단계 완료 여부**: {st.session_state.get('current_step').get('completed', False)}")

    # 간단 검색 시스템 초기화
    try:
        from utils_pdf import initialize_vector_system
        initialize_vector_system()
    except Exception as e:
        st.warning(f"⚠️ 검색 시스템 초기화 실패: {e}")
        st.info("ℹ️ 기본 검색 모드로 진행합니다.")

    # 1) 실행 대상 단계 목록 구성
    current_steps = st.session_state.get('workflow_steps', [])
    
    if not current_steps:
        st.warning("분석할 단계가 없습니다.")
        return

    # 2) prompt_loader에서 해당 단계들 매칭
    try:
        # 프롬프트 블록 로드
        from dsl_to_prompt import load_prompt_blocks
        blocks = load_prompt_blocks()
        extra_blocks = blocks.get("extra", [])
        blocks_by_id = {b["id"]: b for b in extra_blocks}

        # ordered_blocks 대신 workflow_steps 사용
        st.session_state.ordered_blocks = current_steps
        
    except Exception as e:
        st.error(f"❌ 프롬프트 블록 로드 실패: {e}")
        return

    # 3) 진행 표시
    current_step_index = st.session_state.get('current_step_index', 0)
    total_steps = len(current_steps)

    if total_steps == 0:
        st.warning("⚠️ 실행할 분석 단계가 없습니다.")
        return

    progress_percentage = ((current_step_index + 1) / total_steps) * 100
    st.progress(progress_percentage / 100)
    st.write(f"**진행 상황**: {current_step_index + 1} / {total_steps}")

    # 4) 현재 단계 표시 및 실행
    if current_step_index < len(current_steps):
        current_step = current_steps[current_step_index]
        
        # 현재 단계에 해당하는 블록 찾기
        current_block = None
        if current_step.id in blocks_by_id:
            current_block = blocks_by_id[current_step.id]
        
        # 현재 단계의 분석 상태 확인 (수정된 로직)
        step_completed = False
        if current_block:
            # current_block['title']로 저장된 경우 체크
            step_completed = any(h['step'] == current_block['title'] for h in st.session_state.get('cot_history', []))
        else:
            # current_step.title로 저장된 경우 체크
            step_completed = any(h['step'] == current_step.title for h in st.session_state.get('cot_history', []))
        
        # 웹 검색 설정 초기화
        if 'web_search_settings' not in st.session_state:
            st.session_state.web_search_settings = {}
        
        # 현재 단계 표시 및 웹 검색 체크박스
        col1, col2 = st.columns([4, 1])
        
        with col1:
            st.markdown(f"현재 단계: {current_step.title}")
            st.markdown(f"**설명**: {current_step.description}")
            
            # 새로 추가: 블록 구조 정보 표시
            if current_block:
                content_dsl = current_block.get("content_dsl", {})
                
                # templates 정보 표시
                if "templates" in content_dsl:
                    templates = content_dsl["templates"]
                    with st.expander("📋 템플릿 정보", expanded=False):
                        if "tables" in templates:
                            st.write(f"**표 템플릿**: {len(templates['tables'])}개")
                            for table_name, columns in templates["tables"].items():
                                st.write(f"- {table_name}: {len(columns)}개 컬럼")
                        if "analysis_sections" in templates:
                            st.write(f"**분석 섹션**: {len(templates['analysis_sections'])}개")
                        # 새로 추가: alternatives 정보 표시
                        if "alternatives" in templates:
                            st.write(f"**대안 옵션**: {len(templates['alternatives'])}개")
                            for alt in templates["alternatives"]:
                                st.write(f"- {alt.get('name', '대안')}")
                
                # data_contract 정보 표시
                if "data_contract" in content_dsl:
                    contract = content_dsl["data_contract"]
                    with st.expander("📊 데이터 요구사항", expanded=False):
                        if "expected_site_fields" in contract:
                            st.write(f"**필요한 사이트 정보**: {len(contract['expected_site_fields'])}개")
                        if "missing_policy" in contract:
                            st.write(f"**누락 정책**: {contract['missing_policy']}")
                        if "locale_overrides" in contract:
                            st.write(f"**지역 설정**: {len(contract['locale_overrides'])}개 지역")
                
                # analysis_framework.scoring 정보 표시
                framework = content_dsl.get("analysis_framework", {})
                if "scoring" in framework:
                    scoring = framework["scoring"]
                    with st.expander("📈 평가 기준", expanded=False):
                        if "criteria" in scoring:
                            st.write(f"**평가 항목**: {len(scoring['criteria'])}개")
                        if "weights" in scoring:
                            st.write(f"**가중치 설정**: {len(scoring['weights'])}개")
                        if "scale" in scoring:
                            st.write(f"**점수 범위**: {scoring['scale']}")
        
        with col2:
            # 웹 검색 체크박스
            web_search_key = f"web_search_{current_step.id}"
            if web_search_key not in st.session_state.web_search_settings:
                st.session_state.web_search_settings[web_search_key] = False
            
            st.session_state.web_search_settings[web_search_key] = st.checkbox(
                "🌐 웹 검색 포함",
                value=st.session_state.web_search_settings[web_search_key],
                key=web_search_key,
                help="이 단계에서 최신 웹 검색 결과를 포함하여 분석합니다."
            )
        
        # 분석 실행 버튼 (단계가 완료되지 않은 경우에만 표시)
        if not step_completed:
            if current_block:
                button_text = f"{current_block['title']} 분석 실행"
            else:
                button_text = f"{current_step.title} 분석 실행"
            
            if st.button(button_text, type="primary", key=f"analyze_{current_step.id}_{current_step_index}"):
                import contextlib
                with contextlib.suppress(Exception):
                    # current_block이 None인 경우 처리
                    if not current_block:
                        st.error(f"❌ '{current_step.title}' 단계에 해당하는 블록을 찾을 수 없습니다.")
                        st.error(f"단계 ID: {current_step.id}")
                        st.error("사용 가능한 블록들:")
                        for block in extra_blocks:
                            st.error(f"- {block['id']}: {block['title']}")
                        return
                    
                    # PDF 요약 정보 가져오기
                    pdf_summary = get_pdf_summary()
                    if not pdf_summary:
                        st.error("❌ PDF 요약 정보가 없습니다. PDF를 다시 업로드해주세요.")
                        return
                    
                    # 사용자 입력 정보 가져오기
                    user_inputs = get_user_inputs()
                    
                    # 웹 검색 설정 가져오기
                    web_search_key = f"web_search_{current_step.id}"
                    include_web_search = st.session_state.web_search_settings.get(web_search_key, False)
                    
                    # 분석 실행 부분에 디버깅 정보 추가
                    with st.spinner(f"{current_block['title']} 분석 중..."):
                        # DSL을 프롬프트로 변환
                        from dsl_to_prompt import convert_dsl_to_prompt
                        
                        # 이전 분석 결과들 가져오기
                        previous_results = ""
                        if st.session_state.get('cot_history'):
                            previous_results = "\n\n".join([f"**{h['step']}**: {h['result']}" for h in st.session_state.cot_history])
                        
                        # 프롬프트 생성 (웹 검색 설정 반영)
                        prompt = convert_dsl_to_prompt(
                            dsl_block=current_block,
                            user_inputs=user_inputs,
                            previous_summary=previous_results,
                            pdf_summary=pdf_summary,
                            site_fields=st.session_state.get('site_fields', {}),
                            include_web_search=include_web_search  # ✅ 사용자 선택 반영
                        )
                        
                        # 웹 검색 상태 표시
                        if include_web_search:
                            st.info("🌐 웹 검색이 포함된 분석을 실행합니다...")
                        
                        # Claude 분석 실행
                        result = execute_claude_analysis(prompt, current_block['title'])
                        # 실패 가드: 결과가 없거나 실패 메시지면 즉시 중단
                        if not result or result == f"{current_block['title']} 분석 실패":
                            st.error(f"❌ {current_block['title']} 분석 실패")
                            return
                        
                        if result and result != f"{current_block['title']} 분석 실패":
                            # 결과 저장
                            save_step_result(current_step.id, result)
                            append_step_history(current_step.id, current_block['title'], prompt, result)
                            
                            # cot_history에도 추가 (기존 호환성 유지)
                            if 'cot_history' not in st.session_state:
                                st.session_state.cot_history = []
                            st.session_state.cot_history.append({
                                'step': current_block['title'],
                                'result': result
                            })
                            
                            # 자동 저장
                            from user_state import save_user_data
                            save_user_data()
                            
                            st.success(f"✅ {current_block['title']} 분석 완료!")
                            
                            # 분석 완료 후 즉시 결과 표시
                            st.markdown("---")
                            st.markdown(f"### 📋 {current_block['title']} 분석 결과")
                            
                            # 새로 추가: 템플릿 기반 결과 표시
                            content_dsl = current_block.get("content_dsl", {})
                            templates = content_dsl.get("templates", {})
                            
                            if templates and "tables" in templates:
                                # 템플릿이 있는 경우 구조화된 표시
                                output_structure = content_dsl.get("output_structure", [])
                                if output_structure:
                                    parsed_results = parse_analysis_result_by_structure(result, output_structure)
                                    result_tabs = st.tabs(output_structure)
                                    for i, (tab, structure_name) in enumerate(zip(result_tabs, output_structure)):
                                        with tab:
                                            st.markdown(f"### {structure_name}")
                                            content = parsed_results.get(structure_name, "")
                                            if content and not content.startswith("⚠️"):
                                                st.markdown(content)
                                                
                                                # 템플릿 정보 표시
                                                if structure_name in templates.get("tables", {}):
                                                    table_info = templates["tables"][structure_name]
                                                    st.info(f"📋 이 섹션은 {len(table_info)}개 컬럼의 표 형식으로 구성됩니다.")
                                            else:
                                                st.warning("⚠️ 이 구조의 결과를 찾을 수 없습니다.")
                            else:
                                # 기존 방식: 일반 결과 표시
                                output_structure = current_block.get("content_dsl", {}).get("output_structure", [])
                                if output_structure:
                                    parsed_results = parse_analysis_result_by_structure(result, output_structure)
                                    result_tabs = st.tabs(output_structure)
                                    for i, (tab, structure_name) in enumerate(zip(result_tabs, output_structure)):
                                        with tab:
                                            st.markdown(f"### {structure_name}")
                                            content = parsed_results.get(structure_name, "")
                                            if content and not content.startswith("⚠️"):
                                                st.markdown(content)
                                            else:
                                                st.warning("⚠️ 이 구조의 결과를 찾을 수 없습니다.")
                                else:
                                    with st.expander(f"📋 {current_block['title']} - 분석 결과", expanded=True):
                                        st.markdown(result)
                            
                            # 컨트롤 버튼들
                            st.markdown("---")
                            st.markdown("### 분석 제어")
                            col1, col2, col3, col4 = st.columns(4)
                            
                            with col1:
                                if st.button("🔄 다시 분석", key=f"reanalyze_{current_step.id}_{current_step_index}"):
                                    # 재분석 실행
                                    try:
                                        pdf_summary = get_pdf_summary()
                                        user_inputs = get_user_inputs()
                                        
                                        with st.spinner(f"{current_step.title} 재분석 중..."):
                                            from dsl_to_prompt import convert_dsl_to_prompt
                                            
                                            previous_results = ""
                                            if st.session_state.get('cot_history'):
                                                # 현재 단계 결과 제외
                                                previous_results = "\n\n".join([
                                                    f"**{h['step']}**: {h['result']}" 
                                                    for h in st.session_state.cot_history 
                                                    if h['step'] != current_block['title']
                                                ])
                                            
                                            prompt = convert_dsl_to_prompt(
                                                dsl_block=current_block,
                                                user_inputs=user_inputs,
                                                previous_summary=previous_results,
                                                pdf_summary=pdf_summary,
                                                site_fields=st.session_state.get('site_fields', {}),
                                                include_web_search=False
                                            )
                                            
                                            new_result = execute_claude_analysis(prompt, current_block['title'])
                                            
                                            if new_result and new_result != f"{current_block['title']} 분석 실패":
                                                # 기존 결과 업데이트
                                                for h in st.session_state.cot_history:
                                                    if h['step'] == current_block['title']:
                                                        h['result'] = new_result
                                                        break
                                                
                                                save_step_result(current_step.id, new_result)
                                                st.success("✅ 재분석 완료!")
                                                st.rerun()
                                            else:
                                                st.error("❌ 재분석 실패")
                                    except Exception as e:
                                        st.error(f"❌ 재분석 오류: {e}")
                            
                            with col2:
                                if st.button("💬 피드백", key=f"feedback_{current_step.id}_{current_step_index}"):
                                    st.session_state.show_feedback = True
                                    st.rerun()
                            
                            # 피드백 입력 UI
                            if st.session_state.get('show_feedback', False):
                                st.markdown("---")
                                st.markdown("#### 💬 피드백 입력")
                                
                                # 피드백 유형 선택
                                feedback_type = st.selectbox(
                                    "피드백 유형 선택:",
                                    ["추가 분석 요청", "수정 요청", "다른 관점 제시", "구조 변경", "기타"],
                                    key=f"feedback_type_{current_step.id}"
                                )
                                
                                # 피드백 내용 입력
                                feedback_input = st.text_area(
                                    "피드백 내용을 입력해주세요:",
                                    placeholder="개선하고 싶은 부분이나 추가 요청사항을 자유롭게 작성해주세요...",
                                    height=150,
                                    key=f"feedback_input_{current_step.id}"
                                )
                                
                                # 피드백 제출 버튼
                                if st.button("💬 피드백 제출", key=f"submit_feedback_{current_step.id}"):
                                    if feedback_input.strip():
                                        with st.spinner("피드백을 처리하고 있습니다..."):
                                            try:
                                                # 피드백 처리 프롬프트 생성
                                                current_results = st.session_state.current_step_outputs
                                                original_result = current_results.get("original_result", "")
                                                if not original_result:
                                                    # 원본 결과가 없으면 현재 결과를 원본으로 저장
                                                    original_result = "\n\n".join([
                                                        f"**{key}**: {value}" 
                                                        for key, value in current_results.items() 
                                                        if key != "saved" and key != "original_result" and key != "updated_result" and key != "feedback_applied"
                                                    ])
                                                    st.session_state.current_step_outputs["original_result"] = original_result
                                                
                                                feedback_prompt = f"""
기존 분석 결과:
{original_result}

사용자 피드백:
- 유형: {feedback_type}
- 내용: {feedback_input}

위 피드백을 바탕으로 기존 분석 결과를 수정하거나 보완해주세요.
피드백의 의도를 정확히 파악하여 적절한 수정을 제시해주세요.

요청사항:
1. 기존 분석 결과를 바탕으로 사용자의 피드백을 반영한 수정된 분석을 제공해주세요.
2. 피드백 유형에 따라 적절한 수정 방향을 제시해주세요:
   - 추가 분석 요청: 더 자세한 분석이나 새로운 관점 추가
   - 수정 요청: 기존 내용의 오류나 부족한 부분 수정
   - 다른 관점 제시: 새로운 시각이나 접근 방법 제시
   - 구조 변경: 분석 구조나 형식의 변경
   - 기타: 특별한 요청사항에 따른 맞춤형 수정
3. 수정된 결과는 기존 분석의 맥락을 유지하면서 피드백을 반영한 형태로 제공해주세요.
"""
                                                
                                                # 피드백 처리 실행
                                                from agent_executor import execute_agent
                                                updated_result = execute_agent(feedback_prompt)
                                                
                                                # 업데이트된 결과 저장
                                                st.session_state.current_step_outputs["updated_result"] = updated_result
                                                st.session_state.current_step_outputs["feedback_applied"] = True
                                                
                                                # 피드백 히스토리에 추가
                                                if "feedback_history" not in st.session_state:
                                                    st.session_state.feedback_history = []
                                                
                                                st.session_state.feedback_history.append({
                                                    "step": current_step.title,
                                                    "feedback_type": feedback_type,
                                                    "feedback_content": feedback_input,
                                                    "ai_response": updated_result,
                                                    "timestamp": time.time()
                                                })
                                                
                                                # cot_history 업데이트 (마지막 항목을 업데이트된 결과로 교체)
                                                if st.session_state.cot_history:
                                                    st.session_state.cot_history[-1]["result"] = updated_result
                                                
                                                st.success("✅ 피드백이 처리되었습니다!")
                                                st.info(" 피드백이 적용된 결과가 아래에 표시됩니다.")
                                                
                                                # 피드백 적용된 결과 즉시 표시
                                                st.markdown("#### ✨ 피드백 적용된 결과")
                                                st.markdown(updated_result)
                                                
                                            except Exception as e:
                                                st.error(f"❌ 피드백 처리 중 오류가 발생했습니다: {e}")
                                                st.error("오류 상세 정보:")
                                                st.error(f"- 함수: execute_agent")
                                                st.error(f"- 매개변수: {len(feedback_prompt)} 문자")
                                                st.error(f"- 피드백 유형: {feedback_type}")
                                                st.error(f"- 피드백 내용: {feedback_input[:100]}...")
                                                st.error(f"- 오류 타입: {type(e).__name__}")
                                                st.error(f"- 오류 위치: {e.__traceback__.tb_lineno if hasattr(e, '__traceback__') else 'N/A'}")
                                    else:
                                        st.warning("⚠️ 피드백 내용을 입력해주세요.")
                                
                                # 피드백 히스토리 표시
                                if st.session_state.get("feedback_history"):
                                    st.markdown("#### 피드백 히스토리")
                                    for i, feedback in enumerate(st.session_state.feedback_history[-3:], 1):  # 최근 3개만 표시
                                        with st.expander(f"피드백 {i}: {feedback['feedback_type']}", expanded=False):
                                            st.markdown(f"**피드백**: {feedback['feedback_content']}")
                                            st.markdown(f"**AI 응답**: {feedback['ai_response'][:300]}...")
                            
                                # 다시 분석 버튼 (원본 결과로 되돌리기)
                                if st.session_state.get('current_step_outputs', {}).get("feedback_applied"):
                                    if st.button("🔄 원본 결과로 되돌리기", key=f"revert_original_{current_step.id}"):
                                        st.session_state.current_step_outputs["feedback_applied"] = False
                                        if st.session_state.cot_history:
                                            # 원본 결과로 되돌리기
                                            original_result = st.session_state.current_step_outputs.get("original_result", "")
                                            if original_result:
                                                    st.session_state.cot_history[-1]["result"] = original_result
                                        st.rerun()
                            
                                # 피드백 취소 버튼
                                if st.button("❌ 피드백 취소", key=f"cancel_feedback_{current_step.id}"):
                                    st.session_state.show_feedback = False
                                    st.rerun()
                            
                            with col3:
                                if current_step_index > 0:
                                    if st.button("⬅️ 이전 단계", key=f"prev_{current_step.id}_{current_step_index}"):
                                        st.session_state.current_step_index = current_step_index - 1
                                        st.rerun()
                                else:
                                    st.markdown("⬅️ 이전 단계")
                            
                            with col4:
                                if current_step_index < len(current_steps) - 1:
                                    if st.button("➡️ 다음 단계", key=f"next_{current_step.id}_{current_step_index}"):
                                        st.session_state.current_step_index = current_step_index + 1
                                        st.rerun()
                                else:
                                    if st.button("🏁 완료", key=f"finish_{current_step.id}_{current_step_index}"):
                                        st.success("모든 분석이 완료되었습니다!")
                                        st.session_state.analysis_completed = True
                        
                        # 다음 단계 안내
                        if current_step_index < len(current_steps) - 1:
                            next_step = current_steps[current_step_index + 1]
                            st.info(f"➡️ 다음 단계: {next_step.title}")
                
            
        
        # 이미 완료된 단계인 경우 결과 표시
        if step_completed:
            st.success(f"✅ {current_step.title} - 분석 완료")
            
            # 결과 표시 - output_structure 기반 탭으로 변경
            step_title_for_history = current_block['title'] if current_block else current_step.title
            step_result = next((h['result'] for h in st.session_state.cot_history if h['step'] == step_title_for_history), "")
            
            # DSL에서 output_structure 가져오기
            output_structure = current_block.get("content_dsl", {}).get("output_structure", []) if current_block else []
            
            if output_structure:
                # 디버깅 정보 표시 (개발 모드)
                if st.session_state.get('debug_mode', False):
                    with st.expander("디버깅 정보", expanded=False):
                        st.markdown("**Output Structure:**")
                        for i, structure in enumerate(output_structure, 1):
                            st.markdown(f"{i}. {structure}")
                        
                        st.markdown("**AI 응답 길이:**")
                        st.markdown(f"{len(step_result)} 문자")
                        
                        st.markdown("**구조명 매칭 결과:**")
                        for structure in output_structure:
                            found = structure in step_result
                            st.markdown(f"- {structure}: {'✅' if found else '❌'}")
                        
                        st.markdown("**전체 응답 미리보기:**")
                        st.code(step_result[:1000] + "..." if len(step_result) > 1000 else step_result)
                
                # 결과를 구조별로 파싱
                parsed_results = parse_analysis_result_by_structure(step_result, output_structure)
                
                # output_structure 기반 탭 생성
                result_tabs = st.tabs(output_structure)
                
                for i, (tab, structure_name) in enumerate(zip(result_tabs, output_structure)):
                    with tab:
                        st.markdown(f"### {structure_name}")
                        
                        content = parsed_results.get(structure_name, "")
                        
                        if content and not content.startswith("⚠️"):
                            st.markdown(content)
                        else:
                            st.warning("⚠️ 이 구조의 결과를 찾을 수 없습니다.")
                            
                            # 디버깅 정보 표시
                            with st.expander("디버깅 정보", expanded=False):
                                st.markdown("**전체 AI 응답:**")
                                st.code(step_result[:1000] + "..." if len(step_result) > 1000 else step_result)
                                
                                st.markdown("**구조명 검색 결과:**")
                                found = structure_name in step_result
                                st.markdown(f"- '{structure_name}' 포함 여부: {'✅' if found else '❌'}")
                                
                                if found:
                                    st.markdown("**관련 부분:**")
                                    idx = step_result.find(structure_name)
                                    context = step_result[max(0, idx-100):idx+len(structure_name)+200]
                                    st.code(context)
            else:
                # output_structure가 없는 경우
                if current_block:
                    expander_title = f"{current_block['title']} - 분석 결과"
                else:
                    expander_title = f"{current_step.title} - 분석 결과"
                
                with st.expander(expander_title, expanded=True):
                    st.markdown(step_result)
            
            # 컨트롤 버튼들 (이미 완료된 단계용)
            st.markdown("---")
            st.markdown("### 분석 제어")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("🔄 다시 분석", key=f"reanalyze_completed_{current_step.id}_{current_step_index}"):
                    # 재분석 실행
                    try:
                        pdf_summary = get_pdf_summary()
                        user_inputs = get_user_inputs()
                        
                        with st.spinner(f"{current_block['title']} 재분석 중..."):
                            from dsl_to_prompt import convert_dsl_to_prompt
                            
                            previous_results = ""
                            if st.session_state.get('cot_history'):
                                # 현재 단계 결과 제외
                                previous_results = "\n\n".join([
                                    f"**{h['step']}**: {h['result']}" 
                                    for h in st.session_state.cot_history 
                                    if h['step'] != current_block['title']
                                ])
                            
                            prompt = convert_dsl_to_prompt(
                                dsl_block=current_block,
                                user_inputs=user_inputs,
                                previous_summary=previous_results,
                                pdf_summary=pdf_summary,
                                site_fields=st.session_state.get('site_fields', {}),
                                include_web_search=False
                            )
                            
                            new_result = execute_claude_analysis(prompt, current_block['title'])
                            
                            if new_result and new_result != f"{current_block['title']} 분석 실패":
                                # 기존 결과 업데이트
                                for h in st.session_state.cot_history:
                                    if h['step'] == current_block['title']:
                                        h['result'] = new_result
                                        break
                                
                                save_step_result(current_step.id, new_result)
                                st.success("✅ 재분석 완료!")
                                st.rerun()
                            else:
                                st.error("❌ 재분석 실패")
                    except Exception as e:
                        st.error(f"❌ 재분석 오류: {e}")
            
            with col2:
                if st.button("💬 피드백", key=f"feedback_completed_{current_step.id}_{current_step_index}"):
                    st.session_state.show_feedback = True
                    st.rerun()
            
            # 완료된 단계에 대한 피드백 입력 UI
            if st.session_state.get('show_feedback', False):
                st.markdown("---")
                st.markdown("#### 💬 피드백 입력")
                
                # 피드백 유형 선택
                feedback_type = st.selectbox(
                    "피드백 유형 선택:",
                    ["추가 분석 요청", "수정 요청", "다른 관점 제시", "구조 변경", "기타"],
                    key=f"feedback_type_completed_{current_step.id}"
                )
                
                # 피드백 내용 입력
                feedback_input = st.text_area(
                    "피드백 내용을 입력해주세요:",
                    placeholder="개선하고 싶은 부분이나 추가 요청사항을 자유롭게 작성해주세요...",
                    height=150,
                    key=f"feedback_input_completed_{current_step.id}"
                )
                
                # 피드백 제출 버튼
                if st.button("💬 피드백 제출", key=f"submit_feedback_completed_{current_step.id}"):
                    if feedback_input.strip():
                        with st.spinner("피드백을 처리하고 있습니다..."):
                            try:
                                # 피드백 처리 프롬프트 생성
                                step_result = next((h['result'] for h in st.session_state.cot_history if h['step'] == current_step.title), "")
                                
                                feedback_prompt = f"""
기존 분석 결과:
{step_result}

사용자 피드백:
- 유형: {feedback_type}
- 내용: {feedback_input}

위 피드백을 바탕으로 기존 분석 결과를 수정하거나 보완해주세요.
피드백의 의도를 정확히 파악하여 적절한 수정을 제시해주세요.

요청사항:
1. 기존 분석 결과를 바탕으로 사용자의 피드백을 반영한 수정된 분석을 제공해주세요.
2. 피드백 유형에 따라 적절한 수정 방향을 제시해주세요:
   - 추가 분석 요청: 더 자세한 분석이나 새로운 관점 추가
   - 수정 요청: 기존 내용의 오류나 부족한 부분 수정
   - 다른 관점 제시: 새로운 시각이나 접근 방법 제시
   - 구조 변경: 분석 구조나 형식의 변경
   - 기타: 특별한 요청사항에 따른 맞춤형 수정
3. 수정된 결과는 기존 분석의 맥락을 유지하면서 피드백을 반영한 형태로 제공해주세요.
"""
                                
                                # 피드백 처리 실행
                                from agent_executor import execute_agent
                                updated_result = execute_agent(feedback_prompt)
                                
                                # cot_history 업데이트
                                for h in st.session_state.cot_history:
                                    if h['step'] == current_step.title:
                                        h['result'] = updated_result
                                        break
                                
                                # 피드백 히스토리에 추가
                                if "feedback_history" not in st.session_state:
                                    st.session_state.feedback_history = []
                                
                                st.session_state.feedback_history.append({
                                    "step": current_step.title,
                                    "feedback_type": feedback_type,
                                    "feedback_content": feedback_input,
                                    "ai_response": updated_result,
                                    "timestamp": time.time()
                                })
                                
                                st.success("✅ 피드백이 처리되었습니다!")
                                st.info(" 피드백이 적용된 결과가 아래에 표시됩니다.")
                                
                                # 피드백 적용된 결과 즉시 표시
                                st.markdown("#### ✨ 피드백 적용된 결과")
                                st.markdown(updated_result)
                                
                            except Exception as e:
                                st.error(f"❌ 피드백 처리 중 오류가 발생했습니다: {e}")
                    else:
                        st.warning("⚠️ 피드백 내용을 입력해주세요.")
                
                # 피드백 히스토리 표시
                if st.session_state.get("feedback_history"):
                    st.markdown("#### 피드백 히스토리")
                    for i, feedback in enumerate(st.session_state.feedback_history[-3:], 1):
                        with st.expander(f"피드백 {i}: {feedback['feedback_type']}", expanded=False):
                            st.markdown(f"**피드백**: {feedback['feedback_content']}")
                            st.markdown(f"**AI 응답**: {feedback['ai_response'][:300]}...")
                
                # 피드백 취소 버튼
                if st.button("❌ 피드백 취소", key=f"cancel_feedback_completed_{current_step.id}"):
                    st.session_state.show_feedback = False
                    st.rerun()
            
            with col3:
                if current_step_index > 0:
                    if st.button("⬅️ 이전 단계", key=f"prev_completed_{current_step.id}_{current_step_index}"):
                        st.session_state.current_step_index = current_step_index - 1
                        st.rerun()
                else:
                    st.markdown("⬅️ 이전 단계")
            
            with col4:
                if current_step_index < len(current_steps) - 1:
                    if st.button("➡️ 다음 단계", key=f"next_completed_{current_step.id}_{current_step_index}"):
                        st.session_state.current_step_index = current_step_index + 1
                        st.rerun()
                else:
                    if st.button("🏁 완료", key=f"finish_completed_{current_step.id}_{current_step_index}"):
                        st.success("모든 분석이 완료되었습니다!")
                        st.session_state.analysis_completed = True
            
            # 다음 단계 안내
            if current_step_index < len(current_steps) - 1:
                next_step = current_steps[current_step_index + 1]
                st.info(f"➡️ 다음 단계: {next_step.title}")

def render_optimization_tab():
    """최적화 조건 탭 렌더링"""
    st.header("최적화 조건 분석")
    
    # 분석 결과 확인 (선택적)
    has_analysis = st.session_state.get('cot_history')
    if not has_analysis:
        st.info("💡 분석 결과가 없어도 외부 문서를 업로드하여 최적화 조건을 분석할 수 있습니다.")
    
    st.info(" 기존 분석 결과를 바탕으로 매스별 최적화 조건을 자동으로 분석합니다.")
    
    # 외부 문서 업로드 (선택사항)
    st.subheader("외부 분석 문서 업로드 (선택사항)")
    
    # 드롭다운으로 변경
    upload_option = st.selectbox(
        "분석 문서 업로드 방식",
        ["문서 업로드", "텍스트 직접 입력", "사용하지 않음"],
        help="💡 미리 분석된 PDF나 Word 문서를 업로드하거나 텍스트를 직접 입력하여 최적화 조건 분석에 활용합니다.",
        key="optimization_upload_option"
    )
    
    external_analysis_content = ""
    
    if upload_option == "문서 업로드":
        uploaded_file = st.file_uploader(
            "분석 문서 업로드 (PDF, DOCX, DOC)",
            type=['pdf', 'docx', 'doc'],
            help="미리 분석된 문서를 업로드하면 해당 내용을 최적화 조건 분석에 활용합니다.",
            key="optimization_file_uploader"
        )
        
        # 업로드된 파일 처리
        if uploaded_file is not None:
            try:
                with st.spinner("문서를 분석하고 있습니다..."):
                    if uploaded_file.type == "application/pdf":
                        from utils_pdf import extract_text_from_pdf
                        import tempfile
                        import os
                        
                        # 임시 파일로 저장
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                            tmp_file.write(uploaded_file.getvalue())
                            tmp_file_path = tmp_file.name
                        
                        # PDF 텍스트 추출
                        pdf_text = extract_text_from_pdf(tmp_file_path)
                        os.unlink(tmp_file_path)  # 임시 파일 삭제
                        
                        # 텍스트 요약
                        if pdf_text:
                            # 구조화된 형식이 제거된 깨끗한 텍스트 사용
                            external_analysis_content = f"**업로드된 PDF 분석 내용:**\n{pdf_text[:2000]}..." if len(pdf_text) > 2000 else f"**업로드된 PDF 분석 내용:**\n{pdf_text}"
                            st.success(f"✅ PDF 문서 분석 완료! (총 {len(pdf_text)}자)")
                        else:
                            st.warning("⚠️ PDF에서 텍스트를 추출할 수 없습니다.")
                            
                    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                        import docx
                        import io
                        
                        try:
                            # DOCX 파일 처리
                            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                            else:
                                # DOC 파일은 간단한 텍스트 추출 (실제로는 더 복잡한 처리가 필요할 수 있음)
                                st.warning("⚠️ DOC 파일 형식은 제한적으로 지원됩니다. DOCX 형식을 권장합니다.")
                                return
                            
                            # 텍스트 추출
                            doc_text = ""
                            for paragraph in doc.paragraphs:
                                doc_text += paragraph.text + "\n"
                            
                            if doc_text.strip():
                                external_analysis_content = f"**업로드된 Word 문서 분석 내용:**\n{doc_text[:2000]}..." if len(doc_text) > 2000 else f"**업로드된 Word 문서 분석 내용:**\n{doc_text}"
                                st.success(f"✅ Word 문서 분석 완료! (총 {len(doc_text)}자)")
                            else:
                                st.warning("⚠️ Word 문서에서 텍스트를 추출할 수 없습니다.")
                                
                        except Exception as e:
                            st.error(f"❌ Word 문서 처리 중 오류: {e}")
                            
            except Exception as e:
                st.error(f"❌ 문서 처리 중 오류가 발생했습니다: {e}")
    
    elif upload_option == "텍스트 직접 입력":
        external_analysis_content = st.text_area(
            "분석 내용 직접 입력",
            placeholder="분석된 내용을 직접 입력하세요. 이 내용이 최적화 조건 분석에 활용됩니다.",
            height=200,
            help="💡 분석된 내용을 직접 입력하여 최적화 조건 분석에 활용합니다.",
            key="optimization_text_input"
        )
        if external_analysis_content:
            external_analysis_content = f"**직접 입력된 분석 내용:**\n{external_analysis_content}"
    
    # 업로드된 문서 내용 미리보기
    if external_analysis_content:
        with st.expander("📄 업로드된 문서 내용 미리보기"):
            st.markdown(external_analysis_content)
    
    # 자동 분석 실행
    if st.button("매스별 최적화 조건 자동 분석", type="primary"):
        with st.spinner("매스별 최적화 조건을 자동으로 분석하고 있습니다..."):
            try:
                # 사용자 입력 가져오기
                from user_state import get_user_inputs
                user_inputs = get_user_inputs()
                
                # 분석 결과 요약 (내부 분석 + 외부 문서)
                if has_analysis:
                    internal_analysis = "\n\n".join([
                        f"**{h['step']}**: {h.get('result', '')}"
                        for h in st.session_state.cot_history
                    ])
                else:
                    internal_analysis = ""
                
                # 외부 문서 내용과 내부 분석을 결합
                if external_analysis_content and internal_analysis:
                    analysis_summary = f"{external_analysis_content}\n\n**내부 분석 결과:**\n{internal_analysis}"
                elif external_analysis_content:
                    analysis_summary = external_analysis_content
                elif internal_analysis:
                    analysis_summary = internal_analysis
                else:
                    analysis_summary = "분석 결과가 없습니다. 프로젝트 정보만을 기반으로 최적화 조건을 분석합니다."
                
                # 새로운 생성 함수 사용 (분석 결과가 없으면 빈 리스트 전달)
                cot_history = st.session_state.cot_history if has_analysis else []
                
                # 외부 문서 내용을 포함한 최적화 분석 실행
                optimization_result = generate_optimization_analysis_with_external_content(user_inputs, cot_history, analysis_summary)
                
                # 결과를 session_state에 저장
                st.session_state.optimization_result = optimization_result
                
                # 결과 표시
                st.success("✅ 매스별 최적화 조건 분석이 완료되었습니다!")
                
                # 결과를 탭으로 표시
                st.markdown("### 매스별 최적화 조건 분석 결과")
                st.markdown(optimization_result)
                
                # 결과 다운로드 버튼
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"mass_optimization_conditions_{timestamp}.json"
                
                # JSON 형태로 결과 저장
                optimization_data = {
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "project_info": {
                        "project_name": user_inputs.get('project_name', ''),
                        "building_type": user_inputs.get('building_type', ''),
                        "site_location": user_inputs.get('site_location', ''),
                        "owner": user_inputs.get('owner', ''),
                        "site_area": user_inputs.get('site_area', '')
                    },
                    "analysis_result": optimization_result,
                    "external_content_used": bool(external_analysis_content)
                }
                
                # JSON 파일 다운로드 버튼
                st.download_button(
                    label=" 매스별 최적화 조건 분석 결과 다운로드",
                    data=json.dumps(optimization_data, ensure_ascii=False, indent=2),
                    file_name=filename,
                    mime="application/json"
                )
                
            except Exception as e:
                st.error(f"❌ 매스별 최적화 조건 분석 중 오류가 발생했습니다: {e}")
    
    # 이전 분석 결과 표시
    if st.session_state.get('optimization_result'):
        st.markdown("### 이전 매스별 최적화 조건 분석 결과")
        with st.expander("이전 분석 결과 보기", expanded=False):
            st.markdown(st.session_state.optimization_result)

def render_tabbed_interface():
    """탭 기반 인터페이스 렌더링"""
    # 메인 헤더 수정: dAI+ 메인 (검은색, 큰 글씨), ArchInsight 분석 시스템 작은 글씨 (정확한 아래정렬)
    st.markdown("""
    <div style="display: flex; align-items: flex-end; gap: 10px;">
        <h1 style="margin: 0; padding: 0; font-size: 3.5rem; font-weight: bold; color: #000000; line-height: 1;">dAI+</h1>
        <span style="font-size: 1rem; color: #666; font-weight: normal; margin: 0; padding: 0; line-height: 1; margin-bottom: 12px;">ArchInsight 분석 시스템</span>
    </div>
    """, unsafe_allow_html=True)
    
    # 탭 생성 (최적화 조건 탭 추가)
    tab1, tab2, tab3, tab4, tab5 = st.tabs([ # tab6 제거
        "분석 워크플로우", 
        "분석 결과 보고서 생성", 
        "최적화 조건",
        "건축설계 발표용 Narrative", 
        "ArchiRender",
        # "보고서 생성"
    ])
    
    with tab1:
        render_analysis_workflow()
    
    with tab2:
        render_report_tab()
    
    with tab3:
        render_optimization_tab()
    
    with tab4:
        render_claude_narrative_tab()
    
    with tab5:
        render_midjourney_prompt_tab()
    
    # with tab6:
    #     render_report_generation_tab()

def render_report_tab():
    """분석 결과 탭 렌더링"""
    st.header("분석 결과")
    
    if st.session_state.get('cot_history'):
        st.success("✅ 분석이 완료되었습니다!")
        
        # 각 단계별 결과 표시
        st.subheader("각 단계별 분석 결과")
        for i, history in enumerate(st.session_state.cot_history, 1):
            with st.expander(f"{i}. {history['step']}", expanded=True):
                # 요약과 인사이트 제거
                # st.markdown(f"**요약:** {history.get('summary', '')}")
                # st.markdown(f"**인사이트:** {history.get('insight', '')}")
                # st.markdown("---")
                st.markdown(history.get('result', ''))
        
        # PDF/Word 다운로드 섹션 추가 (Tab 6에서 이동)
        st.markdown("---")
        st.subheader("📄 분석 결과 다운로드")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📄 PDF 보고서")
            if st.button("📄 PDF 다운로드", type="primary", key="pdf_download_analysis"):
                with st.spinner("PDF 보고서를 생성하고 있습니다..."):
                    try:
                        from report_generator import generate_report_content, generate_pdf_report
                        report_content = generate_report_content(
                            "전체 분석 보고서", 
                            True, 
                            True, 
                            False
                        )
                        
                        pdf_data = generate_pdf_report(report_content, st.session_state)
                        st.download_button(
                            label="📄 PDF 다운로드",
                            data=pdf_data,
                            file_name=f"{st.session_state.get('project_name', '분석보고서')}_보고서.pdf",
                            mime="application/pdf",
                            key="pdf_download_analysis_final"
                        )
                        
                    except Exception as e:
                        st.error(f"PDF 생성 오류: {e}")
        
        with col2:
            st.markdown("#### 📄 Word 보고서")
        if st.button("📄 Word 다운로드", type="primary", key="word_download_analysis"):
            with st.spinner("Word 보고서를 생성하고 있습니다..."):
                try:
                    from report_generator import generate_report_content, generate_word_report
                    report_content = generate_report_content(
                        "전체 분석 보고서", 
                        True, 
                        True, 
                        False
                    )
                        
                    word_data = generate_word_report(report_content, st.session_state)
                    st.download_button(
                        label="📄 Word 다운로드",
                        data=word_data,
                        file_name=f"{st.session_state.get('project_name', '분석보고서')}_보고서.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="word_download_analysis_final"
                    )
                    
                except Exception as e:
                    st.error(f"Word 생성 오류: {e}")
        
    else:
        st.info("분석을 먼저 진행해주세요.")

def create_improved_narrative_prompt(project_info, direction_choices, analysis_summary):
    """개선된 Narrative 프롬프트 생성 함수"""
    
    # 프로젝트 정보 포맷팅
    project_info_text = f"""
### 프로젝트 기본 정보
- 프로젝트명: {project_info.get('project_name', 'N/A')}
- 건물 유형: {project_info.get('building_type', 'N/A')}
- 건축주: {project_info.get('owner', 'N/A')}
- 발주처 특성: {project_info.get('owner_type', 'N/A')}
- 대지 위치: {project_info.get('site_location', 'N/A')}
- 대지 면적: {project_info.get('site_area', 'N/A')}
- 건물 규모: {project_info.get('building_scale', 'N/A')}
- 주변 환경: {project_info.get('surrounding_env', 'N/A')}
- 지역적 맥락: {project_info.get('regional_context', 'N/A')}
"""

    # 방향성 선택 포맷팅
    direction_text = f"""
### 선택된 Narrative 방향성
1. **감성/논리 비율**: {direction_choices.get('emotion_logic_ratio', 'N/A')}
2. **서술 스타일**: {direction_choices.get('narrative_tone', 'N/A')}
3. **키 메시지 방향**: {direction_choices.get('key_message_direction', 'N/A')}
4. **건축적 가치 우선순위**: {direction_choices.get('architectural_value', 'N/A')}
5. **내러티브 전개 방식**: {direction_choices.get('narrative_structure', 'N/A')}
6. **강조 설계 요소**: {', '.join(direction_choices.get('design_elements', []))}
"""

    return f"""
# 건축설계 발표용 Narrative 생성 시스템

## 🎯 시스템 개요
이 시스템은 건축설계 발표용 Narrative를 단계별로 생성하는 구조화된 프롬프트 체인입니다. 
입력된 정보를 바탕으로 선택된 방향성에 따라 맞춤형 발표 스토리를 자동 생성합니다.

---

## 📋 입력된 프로젝트 정보

{project_info_text}

## 🎨 선택된 Narrative 방향성

{direction_text}

## 📊 분석 결과 요약

{analysis_summary}

**분석 자료 활용 지침:**
- 위 분석 결과를 바탕으로 각 Part의 내용을 구체적이고 설득력 있게 작성하세요
- 외부 문서의 내용이 있다면 해당 내용을 적절한 Part에 통합하여 활용하세요
- 내부 분석 결과와 외부 문서 내용을 조화롭게 결합하여 일관된 Narrative를 구성하세요

---

## 🏗️ Narrative 자동 생성 지침

**시스템 메시지:** "입력된 정보를 바탕으로 맞춤형 Narrative를 생성합니다."

### 생성 프로세스:
1. **선택된 방향성 분석**: 입력된 선택사항들을 종합 분석
2. **템플릿 커스터마이징**: 기본 템플릿을 선택사항에 맞게 조정  
3. **콘텐츠 자동 생성**: 입력 정보와 선택사항을 바탕으로 각 Part별 내용 생성
4. **일관성 검증**: 전체 Narrative의 논리적 흐름과 일관성 확인
5. **최종 결과물 출력**: 완성된 Narrative 제시

### 반드시 생성해야 할 구조:

```
### Part 1. 📋 프로젝트 기본 정보
프로젝트의 핵심 정보를 간결하고 임팩트 있게 정리하세요. 발주처 특성과 프로젝트 규모를 강조하세요.

### Part 2. 🎯 Core Story: 완벽한 교집합의 발견
선택된 키 메시지 방향에 따라 핵심 스토리를 구성하세요. 감성/논리 비율에 맞는 톤으로 서술하세요.

### Part 3. 📍 땅이 주는 답
Context-Driven 방식으로 대지와 지역적 맥락을 분석하세요. "이 땅에서 발견한 진실" 같은 스토리적 접근을 사용하세요.

### Part 4. 🏢 {project_info.get('owner', '건축주')}이 원하는 미래
Vision 중심으로 발주처의 꿈과 비전을 표현하세요. 발주처 특성에 맞는 미래 시나리오를 제시하세요.

### Part 5. 💡 [컨셉명] 컨셉의 탄생
키워드 기반으로 설계 컨셉의 탄생 과정을 서술하세요. 선택된 건축적 가치 우선순위를 반영하세요.

### Part 6. 🏛️ 교집합이 만든 건축적 해답
선택된 내러티브 전개 방식을 적용하세요. 강조할 설계 요소들을 중심으로 건축적 해답을 제시하세요.

### Part 7. 🎯 Winning Narrative 구성
선택된 서술 스타일과 톤을 완전히 적용하세요. 설득력 있고 감동적인 발표용 톤으로 구성하세요.

### Part 8. 🎯 결론: 완벽한 선택의 이유
모든 선택사항이 만든 완벽한 조합의 이유를 설명하세요. 청중에게 남길 강력한 최종 메시지를 제시하세요.
```

### ⚠️ 필수 지시사항:
1. **구조 준수**: 반드시 8개 Part를 위 순서대로 작성하세요
2. **제목 포함**: 각 Part는 반드시 제목(예: "### Part 1. 📋 프로젝트 기본 정보")과 함께 작성하세요
3. **선택사항 반영**: 모든 선택된 방향성을 해당 Part에 반영하세요
4. **소설적 서술**: 기술적 설명이 아닌 감성적 스토리텔링으로 작성하세요
5. **일관성 유지**: 전체적으로 일관된 톤과 스타일을 유지하세요

### 🎨 작성 스타일 가이드:
- 소설처럼 감성적이고 몰입감 있는 서술
- "이 땅에서 발견한 진실" 같은 스토리적 접근
- 구체적인 공간 경험과 사용자 여정을 생생하게 묘사
- 건축적 해답을 스토리로 풀어내기
- 청중의 감정을 움직이는 서술 방식

### 📝 구체적 작성 예시:
- **감성 중심형**: "햇살이 스며드는 순간, 이 공간은 우리에게 어떤 이야기를 들려줄까요?"
- **비즈니스 중심형**: "이 건축은 단순한 공간을 넘어 브랜드 가치를 극대화하는 전략적 도구입니다"
- **미래지향형**: "2030년, 이 건물은 우리의 라이프스타일을 어떻게 변화시킬까요?"

위 가이드라인을 정확히 따라 8단계 구조화된 Narrative를 작성해주세요.
"""

def validate_narrative_quality(narrative_text):
    """Narrative 품질 검증 함수"""
    issues = []
    
    # 구조 검증
    required_parts = [
        "Part 1", "Part 2", "Part 3", "Part 4", 
        "Part 5", "Part 6", "Part 7", "Part 8"
    ]
    
    for part in required_parts:
        if part not in narrative_text:
            issues.append(f"❌ {part} 누락")
    
    # 길이 검증
    if len(narrative_text) < 2000:
        issues.append("⚠️ Narrative가 너무 짧습니다 (최소 2000자 권장)")
    elif len(narrative_text) > 8000:
        issues.append("⚠️ Narrative가 너무 깁니다 (최대 8000자 권장)")
    
    # 스토리텔링 요소 검증
    storytelling_keywords = ["이 땅", "발견", "경험", "만남", "이야기", "여정"]
    found_keywords = [kw for kw in storytelling_keywords if kw in narrative_text]
    
    if len(found_keywords) < 2:
        issues.append("⚠️ 스토리텔링 요소가 부족합니다")
    
    return issues

def render_claude_narrative_tab():
    """Claude Narrative 탭 렌더링 - 개선된 버전"""
    st.header("건축설계 발표용 Narrative 생성 시스템")
    
    # 분석 결과 확인 - 독립 실행 가능하도록 수정
    has_analysis = st.session_state.get('cot_history')
    if not has_analysis:
        st.info("💡 외부 문서를 업로드하거나 기본 정보만으로도 Narrative를 생성할 수 있습니다.")
    
    st.info(" 건축설계 발표용 Narrative를 단계별로 생성하는 구조화된 시스템입니다.")
    
    # STEP 0: 외부 문서 업로드 (선택사항)
    st.subheader("STEP 0: 외부 분석 문서 업로드 (선택사항)")
    
    # 드롭다운으로 변경
    upload_option = st.selectbox(
        "분석 문서 업로드 방식",
        ["문서 업로드", "텍스트 직접 입력", "사용하지 않음"],
        help="💡 미리 분석된 PDF나 Word 문서를 업로드하거나 텍스트를 직접 입력하여 Narrative 생성에 활용합니다.",
        key="narrative_upload_option"
    )
    
    external_analysis_content = ""
    
    if upload_option == "문서 업로드":
        uploaded_file = st.file_uploader(
            "분석 문서 업로드 (PDF, DOCX, DOC)",
            type=['pdf', 'docx', 'doc'],
            help="미리 분석된 문서를 업로드하면 해당 내용을 Narrative 생성에 활용합니다.",
            key="narrative_file_uploader"
        )
        
        # 업로드된 파일 처리
        if uploaded_file is not None:
            try:
                with st.spinner("문서를 분석하고 있습니다..."):
                    if uploaded_file.type == "application/pdf":
                        from utils_pdf import extract_text_from_pdf
                        import tempfile
                        import os
                        
                        # 임시 파일로 저장
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                            tmp_file.write(uploaded_file.getvalue())
                            tmp_file_path = tmp_file.name
                        
                        # PDF 텍스트 추출
                        pdf_text = extract_text_from_pdf(tmp_file_path)
                        os.unlink(tmp_file_path)  # 임시 파일 삭제
                        
                        # 텍스트 요약
                        if pdf_text:
                            # 구조화된 형식이 제거된 깨끗한 텍스트 사용
                            external_analysis_content = f"**업로드된 PDF 분석 내용:**\n{pdf_text[:2000]}..." if len(pdf_text) > 2000 else f"**업로드된 PDF 분석 내용:**\n{pdf_text}"
                            st.success(f"✅ PDF 문서 분석 완료! (총 {len(pdf_text)}자)")
                        else:
                            st.warning("⚠️ PDF에서 텍스트를 추출할 수 없습니다.")
                            
                    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                        import docx
                        import io
                        
                        try:
                            # DOCX 파일 처리
                            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                            else:
                                # DOC 파일은 간단한 텍스트 추출 (실제로는 더 복잡한 처리가 필요할 수 있음)
                                st.warning("⚠️ DOC 파일 형식은 제한적으로 지원됩니다. DOCX 형식을 권장합니다.")
                                return

                            # 텍스트 추출
                            doc_text = ""
                            for paragraph in doc.paragraphs:
                                doc_text += paragraph.text + "\n"
                            
                            if doc_text.strip():
                                external_analysis_content = f"**업로드된 Word 문서 분석 내용:**\n{doc_text[:2000]}..." if len(doc_text) > 2000 else f"**업로드된 Word 문서 분석 내용:**\n{doc_text}"
                                st.success(f"✅ Word 문서 분석 완료! (총 {len(doc_text)}자)")
                            else:
                                st.warning("⚠️ Word 문서에서 텍스트를 추출할 수 없습니다.")
                                
                        except Exception as e:
                            st.error(f"❌ Word 문서 처리 중 오류: {e}")
                            return
                            
            except Exception as e:
                st.error(f"❌ 문서 처리 중 오류가 발생했습니다: {e}")
    
    elif upload_option == "텍스트 직접 입력":
        external_analysis_content = st.text_area(
            "분석 내용 직접 입력",
            placeholder="분석된 내용을 직접 입력하세요. 이 내용이 Narrative 생성에 활용됩니다.",
            height=200,
            help="💡 분석된 내용을 직접 입력하여 Narrative 생성에 활용합니다.",
            key="narrative_text_input"
        )
        if external_analysis_content:
            external_analysis_content = f"**직접 입력된 분석 내용:**\n{external_analysis_content}"
    
    # 업로드된 문서 내용 미리보기
    if external_analysis_content:
        with st.expander("📄 업로드된 문서 내용 미리보기"):
            st.markdown(external_analysis_content)
    
    # STEP 1: 기본 정보 입력
    st.subheader("STEP 1: 기본 정보 입력")
    col1, col2 = st.columns(2)
    
    with col1:
        project_name = st.text_input("프로젝트명", value=st.session_state.get('project_name', ''))
        building_type = st.text_input("건물 유형", value=st.session_state.get('building_type', ''))
        site_location = st.text_input("대지 위치", value=st.session_state.get('site_location', ''))
        owner = st.text_input("건축주", value=st.session_state.get('owner', ''))
        owner_type = st.selectbox("발주처 특성", ["공공기관", "민간기업", "개인", "교육기관", "의료기관", "문화기관"])
            
    with col2:
        site_area = st.text_input("대지 면적", value=st.session_state.get('site_area', ''))
        building_scale = st.text_input("건물 규모", placeholder="연면적, 층수 등")
        surrounding_env = st.text_area("주변 환경", placeholder="자연환경, 도시환경, 교통, 랜드마크 등")
        regional_context = st.text_area("지역적 맥락", placeholder="역사, 문화, 사회적 특성")
    
    # STEP 2: Narrative 방향 설정
    st.subheader("STEP 2: Narrative 방향 설정")
    
    # 2-1. 감성 ↔ 논리 비율 선택
    st.markdown("#### 2-1. 감성 ↔ 논리 비율 선택")
    emotion_logic_ratio = st.selectbox(
        "감성/논리 비율을 선택하세요:",
        [
            "A. 감성 중심형 (감성 90% / 논리 10%) - 감정적 울림, 서정적 표현, 상징성 중심",
            "B. 균형형 (감성 60% / 논리 40%) - 사용자 경험 중심 + 분석 기반 논리 서술의 조화",
            "C. 전략 중심형 (감성 30% / 논리 70%) - 기능적 해법 + 분석 데이터 기반 논리 중심",
            "D. 데이터 기반형 (감성 10% / 논리 90%) - 통계·규범·정책 중심 논리적 설득"
        ]
    )
    
    # 2-2. 서술 스타일/톤 선택
    st.markdown("#### 2-2. 서술 스타일/톤 선택")
    narrative_tone = st.selectbox(
        "서술 스타일을 선택하세요:",
        [
            "A. 공공적/진정성형 - 지역사회 기여, 지속가능성, 공동체 가치 강조",
            "B. 비즈니스 중심형 - 경제성, 차별화 전략, 고객 경험 중심 강조",
            "C. 미래지향/비전형 - 변화 주도, 혁신, 미래 라이프스타일 제안",
            "D. 문화/상징성형 - 장소성, 역사 해석, 상징적 메시지 중심",
            "E. 사용자 감성형 - 일상 경험과 공간의 만남, 감각 중심"
        ]
    )
    
    # 2-3. 키 메시지 중심 방향 선택
    st.markdown("#### 2-3. 키 메시지 중심 방향 선택")
    key_message_direction = st.selectbox(
        "핵심 메시지 방향을 선택하세요:",
        [
            "A. Vision 중심형 - 이 건축이 실현할 미래를 제시하는 선언적 서술",
            "B. Problem-Solution형 - 이 문제가 있었고, 이렇게 해결했다는 설계 전략 중심",
            "C. User Journey형 - 사용자의 여정은 어떻게 변화하는가? 사용자 감정·동선 중심",
            "D. Context-Driven형 - 이 땅, 이 장소에서의 필연성은? Site 중심 서술",
            "E. Symbolic Message형 - 이 건물은 어떤 메시지를 담고 있는가? 감정적 울림 강조"
        ]
    )
    
    # 2-4. 건축적 가치 우선순위 선택
    st.markdown("#### 2-4. 건축적 가치 우선순위 선택")
    architectural_value = st.selectbox(
        "건축적 가치 우선순위를 선택하세요:",
        [
            "A. 장소성 우선 - Site-specific한 고유성 추구, 맥락적 건축",
            "B. 기능성 우선 - 사용자 니즈와 효율성 중심, 합리적 건축",
            "C. 미학성 우선 - 아름다움과 감동 추구, 조형적 건축",
            "D. 지속성 우선 - 환경과 미래 세대 고려, 친환경 건축",
            "E. 사회성 우선 - 공동체와 소통 중심, 공공적 건축",
            "F. 혁신성 우선 - 새로운 가능성 탐구, 실험적 건축"
        ]
    )
    
    # 2-5. 건축적 내러티브 전개 방식 선택
    st.markdown("#### 2-5. 건축적 내러티브 전개 방식 선택")
    narrative_structure = st.selectbox(
        "내러티브 전개 방식을 선택하세요:",
        [
            "A. 형태 생성 과정형 - 이 형태는 어떻게 탄생했는가? 대지→매스→공간→디테일 순차 전개",
            "B. 공간 경험 여정형 - 사용자는 어떤 공간을 경험하는가? 진입→이동→머무름→떠남의 시퀀스",
            "C. 기능 조직 논리형 - 프로그램들이 어떻게 조직되는가? 기능분석→배치전략→공간구성",
            "D. 구조 시스템형 - 건물은 어떤 원리로 서 있는가? 구조체→공간→형태의 통합적 설명",
            "E. 환경 대응 전략형 - 자연과 건축이 어떻게 만나는가? 미기후→배치→형태→재료 연계",
            "F. 문화적 해석형 - 전통과 현대가 어떻게 만나는가? 역사적 맥락→현대적 번역→공간화"
        ]
    )
    
    # 2-6. 강조할 설계 요소 선택 (복수 선택 가능)
    st.markdown("#### 2-6. 강조할 설계 요소 선택 (복수 선택 가능)")
    design_elements = st.multiselect(
        "강조할 설계 요소를 선택하세요:",
        [
            "매스/형태 - 조형적 아름다움, 상징성으로 시각적 임팩트",
            "공간 구성 - 동선, 기능 배치의 합리성으로 사용성 어필",
            "친환경/지속가능 - 에너지 효율, 친환경 기술로 사회적 가치",
            "기술/혁신 - 신기술 적용, 스마트 시스템으로 선진성 강조",
            "경제성 - 건설비, 운영비 절감으로 실용성 어필",
            "안전성 - 구조적 안정, 방재 계획으로 신뢰성 구축",
            "문화/역사 - 지역성, 전통의 현대적 해석으로 정체성 강화",
            "사용자 경험 - 편의성, 접근성, 쾌적성으로 만족도 제고"
        ]
    )
    
    # 모델 선택 옵션 추가
    st.markdown("#### 모델 선택 (고급 옵션)")
    selected_model = st.selectbox(
        "Narrative 생성에 사용할 모델을 선택하세요:",
        [
            "claude-sonnet-4-20250514 (권장 - 최고 품질)",
            "claude-opus-4-1-20250805 (최신 - 고성능)",
            "claude-3-5-sonnet-20241022 (기본 - 안정적)"
        ],
        index=0
    )

    # 선택된 모델 추출
    model_mapping = {
        "claude-sonnet-4-20250514 (권장 - 최고 품질)": "claude-sonnet-4-20250514",
        "claude-opus-4-1-20250805 (최신 - 고성능)": "claude-opus-4-1-20250805",
        "claude-3-5-sonnet-20241022 (기본 - 안정적)": "claude-3-5-sonnet-20241022"
    }
    selected_model_id = model_mapping[selected_model]
    
    # STEP 3: Narrative 생성
    st.subheader("STEP 3: Narrative 생성")
    if st.button("Narrative 생성", type="primary"):
        if not all([project_name, building_type, owner]):
            st.error("❌ 기본 정보를 모두 입력해주세요.")
            return
        
        with st.spinner("Narrative를 생성하고 있습니다..."):
            try:
                # 분석 결과 요약 (내부 분석 + 외부 문서)
                has_analysis = st.session_state.get('cot_history')
                if has_analysis:
                    internal_analysis = "\n\n".join([
                    f"**{h['step']}**: {h.get('summary', '')}"
                    for h in st.session_state.cot_history
                ])
                else:
                    internal_analysis = ""
                
                # 외부 문서 내용과 내부 분석을 결합
                if external_analysis_content and internal_analysis:
                    analysis_summary = f"{external_analysis_content}\n\n**내부 분석 결과:**\n{internal_analysis}"
                elif external_analysis_content:
                    analysis_summary = external_analysis_content
                elif internal_analysis:
                    analysis_summary = internal_analysis
                else:
                    analysis_summary = "분석 결과가 없습니다. 프로젝트 정보만을 기반으로 Narrative를 생성합니다."
                
                # 프로젝트 정보 수집
                project_info = {
                    'project_name': project_name,
                    'building_type': building_type,
                    'owner': owner,
                    'owner_type': owner_type,
                    'site_location': site_location,
                    'site_area': site_area,
                    'building_scale': building_scale,
                    'surrounding_env': surrounding_env,
                    'regional_context': regional_context
                }
                
                # 방향성 선택 수집
                direction_choices = {
                    'emotion_logic_ratio': emotion_logic_ratio,
                    'narrative_tone': narrative_tone,
                    'key_message_direction': key_message_direction,
                    'architectural_value': architectural_value,
                    'narrative_structure': narrative_structure,
                    'design_elements': design_elements
                }
                
                # 개선된 프롬프트 생성
                narrative_prompt = create_improved_narrative_prompt(
                    project_info, direction_choices, analysis_summary
                )
                
                # Narrative 생성 함수 호출
                from agent_executor import generate_narrative
                narrative_result = generate_narrative(narrative_prompt)
                
                # 결과를 세션에 저장
                st.session_state.narrative_result = narrative_result
                
                # 결과 표시
                st.success("✅ Narrative 생성 완료!")
                st.markdown("### 📝 생성된 Narrative")
                st.markdown(narrative_result)
                
                # 품질 검증
                quality_issues = validate_narrative_quality(narrative_result)
                if quality_issues:
                    st.warning("⚠️ Narrative 품질 검증 결과:")
                    for issue in quality_issues:
                        st.write(issue)
                else:
                    st.success("✅ Narrative 품질 검증 통과!")
                
                # 다운로드 버튼
                st.download_button(
                    label="📄 Narrative 다운로드",
                    data=narrative_result,
                    file_name=f"{project_name}_Narrative.txt",
                    mime="text/plain"
                )
                
                # 프롬프트 미리보기 (디버깅용)
                with st.expander("🔍 생성된 프롬프트 미리보기"):
                    st.code(narrative_prompt, language="markdown")
                
            except Exception as e:
                st.error(f"❌ Narrative 생성 중 오류가 발생했습니다: {e}")
    
    # STEP 4: 피드백 및 수정
    st.subheader("STEP 4: 피드백 및 수정")
    
    # 생성된 Narrative가 있는지 확인
    if 'narrative_result' in st.session_state:
        st.success("✅ Narrative가 생성되었습니다. 피드백을 통해 수정할 수 있습니다.")
        
        # 피드백 입력
        st.markdown("#### 💬 피드백 입력")
        feedback_type = st.selectbox(
            "피드백 유형:",
            ["전체 톤 조정", "특정 Part 수정", "구조 변경", "내용 추가", "길이 조정", "기타"]
        )
        
        feedback_content = st.text_area(
            "구체적인 피드백 내용:",
            placeholder="예: Part 2를 더 감성적으로 작성해주세요. 또는 전체적으로 더 간결하게 해주세요.",
            height=100
        )
        
        if st.button("🔄 피드백 반영하여 재생성", type="secondary"):
            if feedback_content.strip():
                with st.spinner("피드백을 반영하여 Narrative를 재생성하고 있습니다..."):
                    try:
                        # 피드백 반영 프롬프트 생성
                        feedback_prompt = f"""
기존 Narrative:
{st.session_state.narrative_result}

사용자 피드백:
- 유형: {feedback_type}
- 내용: {feedback_content}

위 피드백을 바탕으로 기존 Narrative를 수정해주세요.

요청사항:
1. 기존 Narrative의 구조와 내용을 유지하면서 피드백을 반영
2. 피드백 유형에 따라 적절한 수정 방향 적용:
   - 전체 톤 조정: 전체적인 서술 톤과 스타일 조정
   - 특정 Part 수정: 해당 Part의 내용과 스타일 수정
   - 구조 변경: 전체적인 구조나 순서 조정
   - 내용 추가: 부족한 내용이나 세부사항 추가
   - 길이 조정: 전체 길이를 늘리거나 줄임
   - 기타: 특별한 요청사항에 따른 맞춤형 수정
3. 수정된 결과는 기존 Narrative의 맥락을 유지하면서 피드백을 반영한 형태로 제공

수정된 Narrative를 8개 Part 구조로 다시 작성해주세요.
"""
                        
                        # 재생성 실행
                        from agent_executor import generate_narrative
                        updated_narrative = generate_narrative(feedback_prompt)
                        
                        # 결과 업데이트
                        st.session_state.narrative_result = updated_narrative
                        
                        st.success("✅ 피드백이 반영된 Narrative가 생성되었습니다!")
                        st.markdown("### 📝 수정된 Narrative")
                        st.markdown(updated_narrative)
                        
                        # 품질 검증
                        quality_issues = validate_narrative_quality(updated_narrative)
                        if quality_issues:
                            st.warning("⚠️ 수정된 Narrative 품질 검증 결과:")
                            for issue in quality_issues:
                                st.write(issue)
                        else:
                            st.success("✅ 수정된 Narrative 품질 검증 통과!")
                        
                        # 다운로드 버튼
                        st.download_button(
                            label="📥 수정된 Narrative 다운로드",
                            data=updated_narrative,
                            file_name=f"{project_name}_Narrative_수정본.txt",
                            mime="text/plain"
                        )
                        
                    except Exception as e:
                        st.error(f"❌ 피드백 반영 중 오류가 발생했습니다: {e}")
            else:
                st.warning("⚠️ 피드백 내용을 입력해주세요.")
    else:
        st.info("🔄 Narrative를 먼저 생성한 후 피드백을 입력할 수 있습니다.")

def render_midjourney_prompt_tab():
    """ArchiRender GPT 탭 렌더링"""
    st.header("ArchiRender GPT")
    
    # 분석 결과 확인 (선택적)
    has_analysis = st.session_state.get('cot_history')
    if not has_analysis:
        st.info("💡 분석 결과가 없어도 외부 문서를 업로드하여 프롬프트를 생성할 수 있습니다.")
    
    st.info("Midjourney 이미지 생성 프롬프트를 생성합니다.")
    
    # 외부 문서 업로드 (선택사항)
    st.subheader("외부 분석 문서 업로드 (선택사항)")
    
    # 드롭다운으로 변경
    upload_option = st.selectbox(
        "분석 문서 업로드 방식",
        ["문서 업로드", "텍스트 직접 입력", "사용하지 않음"],
        help="💡 미리 분석된 PDF나 Word 문서를 업로드하거나 텍스트를 직접 입력하여 이미지 프롬프트 생성에 활용합니다.",
        key="midjourney_upload_option"
    )
    
    external_analysis_content = ""
    
    if upload_option == "문서 업로드":
        uploaded_file = st.file_uploader(
            "분석 문서 업로드 (PDF, DOCX, DOC)",
            type=['pdf', 'docx', 'doc'],
            help="미리 분석된 문서를 업로드하면 해당 내용을 이미지 프롬프트 생성에 활용합니다.",
            key="midjourney_file_uploader"
        )
        
        # 업로드된 파일 처리
        if uploaded_file is not None:
            try:
                with st.spinner("문서를 분석하고 있습니다..."):
                    if uploaded_file.type == "application/pdf":
                        from utils_pdf import extract_text_from_pdf
                        import tempfile
                        import os
                        
                        # 임시 파일로 저장
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                            tmp_file.write(uploaded_file.getvalue())
                            tmp_file_path = tmp_file.name
                        
                        # PDF 텍스트 추출
                        pdf_text = extract_text_from_pdf(tmp_file_path)
                        os.unlink(tmp_file_path)  # 임시 파일 삭제
                        
                        # 텍스트 요약
                        if pdf_text:
                            # 구조화된 형식이 제거된 깨끗한 텍스트 사용
                            external_analysis_content = f"**업로드된 PDF 분석 내용:**\n{pdf_text[:2000]}..." if len(pdf_text) > 2000 else f"**업로드된 PDF 분석 내용:**\n{pdf_text}"
                            st.success(f"✅ PDF 문서 분석 완료! (총 {len(pdf_text)}자)")
                        else:
                            st.warning("⚠️ PDF에서 텍스트를 추출할 수 없습니다.")
                            
                    elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                        import docx
                        import io
                        
                        try:
                            # DOCX 파일 처리
                            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
                            else:
                                # DOC 파일은 간단한 텍스트 추출 (실제로는 더 복잡한 처리가 필요할 수 있음)
                                st.warning("⚠️ DOC 파일 형식은 제한적으로 지원됩니다. DOCX 형식을 권장합니다.")
                                return
                            
                            # 텍스트 추출
                            doc_text = ""
                            for paragraph in doc.paragraphs:
                                doc_text += paragraph.text + "\n"
                            
                            if doc_text.strip():
                                external_analysis_content = f"**업로드된 Word 문서 분석 내용:**\n{doc_text[:2000]}..." if len(doc_text) > 2000 else f"**업로드된 Word 문서 분석 내용:**\n{doc_text}"
                                st.success(f"✅ Word 문서 분석 완료! (총 {len(doc_text)}자)")
                            else:
                                st.warning("⚠️ Word 문서에서 텍스트를 추출할 수 없습니다.")
                                
                        except Exception as e:
                            st.error(f"❌ Word 문서 처리 중 오류: {e}")
                            
            except Exception as e:
                st.error(f"❌ 문서 처리 중 오류가 발생했습니다: {e}")
    
    elif upload_option == "텍스트 직접 입력":
        external_analysis_content = st.text_area(
            "분석 내용 직접 입력",
            placeholder="분석된 내용을 직접 입력하세요. 이 내용이 이미지 프롬프트 생성에 활용됩니다.",
            height=200,
            help="💡 분석된 내용을 직접 입력하여 이미지 프롬프트 생성에 활용합니다.",
            key="midjourney_text_input"
        )
        if external_analysis_content:
            external_analysis_content = f"**직접 입력된 분석 내용:**\n{external_analysis_content}"
    
    # 업로드된 문서 내용 미리보기
    if external_analysis_content:
        with st.expander("📄 업로드된 문서 내용 미리보기"):
            st.markdown(external_analysis_content)
    
    # 이미지 생성 옵션
    st.subheader("이미지 생성 옵션")
    image_type = st.selectbox(
        "이미지 유형",
        ["외관 렌더링", "내부 공간", "마스터플랜", "상세도", "컨셉 이미지", "조감도"]
    )
    
    style_preference = st.multiselect(
        "스타일 선호도",
        ["현대적", "미니멀", "자연친화적", "고급스러운", "기능적", "예술적", "상업적"]
    )
    
    additional_description = st.text_area(
        "추가 설명",
        placeholder="특별히 강조하고 싶은 요소나 스타일을 자유롭게 입력하세요."
    )
    
    # 프롬프트 생성
    if st.button("프롬프트 생성", type="primary"):
        with st.spinner("이미지 생성 프롬프트를 생성하고 있습니다..."):
            try:
                # 사용자 입력 가져오기
                from user_state import get_user_inputs
                user_inputs = get_user_inputs()
                
                # 이미지 설정 구성
                image_settings = {
                    'image_type': image_type,
                    'style_preference': style_preference,
                    'additional_description': additional_description,
                    'external_analysis_content': external_analysis_content
                }
                
                # generate_midjourney_prompt 함수 호출
                prompt_result = generate_midjourney_prompt(
                    user_inputs, 
                    st.session_state.get('cot_history', []), 
                    image_settings
                )
                
                # 결과 표시
                st.success("✅ 이미지 생성 프롬프트 생성 완료!")
                st.markdown("### 생성된 프롬프트")
                st.markdown(prompt_result)
                
                # 다운로드 버튼
                st.download_button(
                    label="📄 프롬프트 다운로드",
                    data=prompt_result,
                    file_name=f"{st.session_state.get('project_name', 'project')}_image_prompt.txt",
                    mime="text/plain"
                )
                
            except Exception as e:
                st.error(f"❌ 프롬프트 생성 중 오류가 발생했습니다: {e}")
    
    # 사용 가이드
    st.subheader("📖 사용 가이드")
    st.markdown("""
    1. **프롬프트 복사**: 생성된 프롬프트를 복사합니다.
    2. **Midjourney 접속**: Discord에서 Midjourney 봇을 찾습니다.
    3. **프롬프트 입력**: `/imagine` 명령어와 함께 프롬프트를 입력합니다.
    4. **이미지 생성**: Midjourney가 이미지를 생성할 때까지 기다립니다.
    5. **결과 확인**: 생성된 이미지를 확인하고 필요시 변형을 요청합니다.
    """)

# def render_report_generation_tab():
#     """보고서 생성 탭 렌더링 - 순서 변경"""
#     st.header("📄 보고서 생성")
    
#     # 분석 결과 확인 (선택적)
#     has_analysis = st.session_state.get('cot_history')
#     if not has_analysis:
#         st.info("💡 분석 결과가 없어도 외부 문서를 업로드하여 보고서를 생성할 수 있습니다.")
    
#     st.success("✅ 분석 결과를 바탕으로 다양한 형태의 보고서를 생성할 수 있습니다.")
    
#     # 사용자 입력 가져오기
#     from user_state import get_user_inputs
#     user_inputs = get_user_inputs()
    
#     # 분석 결과 수집
#     analysis_results = []
#     if st.session_state.get('cot_history'):
#         for i, history in enumerate(st.session_state.cot_history, 1):
#             analysis_results.append({
#                 'step': history.get('step', f'단계 {i}'),
#                 'summary': history.get('summary', ''),
#                 'insight': history.get('insight', ''),
#                 'result': history.get('result', '')
#             })
    
#     # 프로젝트 정보
#     project_info = {
#         'project_name': user_inputs.get('project_name', '프로젝트'),
#         'owner': user_inputs.get('owner', ''),
#         'site_location': user_inputs.get('site_location', ''),
#         'site_area': user_inputs.get('site_area', ''),
#         'building_type': user_inputs.get('building_type', ''),
#         'project_goal': user_inputs.get('project_goal', '')
#     }
    
#     # 1. 문서 보고서 (맨 상단)
#     st.subheader("📄 문서 보고서")
#     report_type = st.selectbox(
#         "보고서 유형",
#         ["전체 분석 보고서", "요약 보고서", "전문가 보고서", "클라이언트 보고서"],
#         key="report_type_generation"
#     )
    
#     include_charts = st.checkbox("차트 포함", value=True, key="charts_generation")
#     include_recommendations = st.checkbox("💡 권장사항 포함", value=True, key="recommendations_generation")
#     include_appendix = st.checkbox("📋 부록 포함", value=False, key="appendix_generation")
    
#     if st.button("📄 보고서 생성", type="primary", key="generate_report_generation"):
#         with st.spinner("보고서를 생성하고 있습니다..."):
#             try:
#                 # 보고서 내용 생성
#                 from report_generator import generate_report_content, generate_pdf_report, generate_word_report
#                 report_content = generate_report_content(
#                     report_type, 
#                     include_charts, 
#                     include_recommendations, 
#                     include_appendix
#                 )
                
#                 # 다운로드 버튼들
#                 col_a, col_b, col_c = st.columns(3)
                
#                 with col_a:
#                     st.download_button(
#                         label="📄 TXT 다운로드",
#                         data=report_content,
#                         file_name=f"{st.session_state.get('project_name', '분석보고서')}_보고서.txt",
#                         mime="text/plain"
#                     )
                
#                 with col_b:
#                     try:
#                         pdf_data = generate_pdf_report(report_content, st.session_state)
#                         st.download_button(
#                             label="📄 PDF 다운로드",
#                             data=pdf_data,
#                             file_name=f"{st.session_state.get('project_name', '분석보고서')}_보고서.pdf",
#                             mime="application/pdf"
#                         )
#                     except Exception as e:
#                         st.error(f"PDF 생성 오류: {e}")
                
#                 with col_c:
#                     try:
#                         word_data = generate_word_report(report_content, st.session_state)
#                         st.download_button(
#                             label="📄 Word 다운로드",
#                             data=word_data,
#                             file_name=f"{st.session_state.get('project_name', '분석보고서')}_보고서.docx",
#                             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                         )
#                     except Exception as e:
#                         st.error(f"Word 생성 오류: {e}")
                
#                 # 보고서 내용 미리보기
#                 st.subheader("📋 보고서 미리보기")
#                 st.markdown(report_content[:2000] + ("..." if len(report_content) > 2000 else ""))
                
#             except Exception as e:
#                 st.error(f"❌ 보고서 생성 중 오류가 발생했습니다: {e}")

#     st.markdown("---")
    
#     # 2. 웹페이지 생성 (중간)
#     st.subheader("📄 웹페이지 생성")
#     from webpage_generator import create_webpage_download_button
#     create_webpage_download_button(analysis_results, project_info, show_warning=False)
    
#     st.markdown("---")
    
#     # 3. 분석 보고서 (맨 하단)
#     st.subheader("분석 보고서")
#     st.markdown("#### 프로젝트 기본 정보")
#     project_info_text = f"""
#     **프로젝트명**: {user_inputs.get('project_name', 'N/A')}
#     **건축주**: {user_inputs.get('owner', 'N/A')}
#     **대지위치**: {user_inputs.get('site_location', 'N/A')}
#     **대지면적**: {user_inputs.get('site_area', 'N/A')}
#     **건물용도**: {user_inputs.get('building_type', 'N/A')}
#     **프로젝트 목표**: {user_inputs.get('project_goal', 'N/A')}
#     """
#     st.markdown(project_info_text)
    
#     # 분석 결과 요약 제거 - 중복되는 부분 삭제

def parse_analysis_result_by_structure(result: str, output_structure: list) -> dict:
    """완전히 개선된 분석 결과 파싱 함수"""
    parsed_results = {}
    
    # 각 구조별로 정확한 경계 찾기
    for i, structure in enumerate(output_structure, 1):
        # 정확한 마커 패턴들 (우선순위 순)
        markers = [
            f"## {i}. {structure}",
            f"## {structure}",
            f"{i}. {structure}",
            f"### {structure}",
            f"**{structure}**",
            structure
        ]
        
        content = None
        start_idx = -1
        used_marker = None
        
        # 마커 찾기
        for marker in markers:
            start_idx = result.find(marker)
            if start_idx != -1:
                used_marker = marker
                break
        
        if start_idx != -1:
            # 다음 구조의 시작 위치 찾기
            end_idx = len(result)
            
            # 다음 번호의 구조 찾기
            for j, next_structure in enumerate(output_structure[i:], i+1):
                next_markers = [
                    f"## {j}. {next_structure}",
                    f"## {next_structure}",
                    f"{j}. {next_structure}",
                    f"### {next_structure}",
                    f"**{next_structure}**"
                ]
                
                for next_marker in next_markers:
                    next_idx = result.find(next_marker, start_idx + len(used_marker))
                    if next_idx != -1 and next_idx < end_idx:
                        end_idx = next_idx
                        break
                
                if end_idx < len(result):
                    break
            
            # 내용 추출
            content_start = start_idx + len(used_marker)
            content = result[content_start:end_idx].strip()
            
            # 앞뒤 공백 및 불필요한 문자 제거
            content = content.strip()
            if content.startswith('\n'):
                content = content[1:]
            if content.endswith('\n'):
                content = content[:-1]
            
            # 내용이 너무 짧으면 무효로 처리
            if len(content) < 10:
                content = None
        
        # 내용을 찾지 못한 경우
        if not content:
            # 키워드 기반 검색
            keywords = structure.lower().split()
            lines = result.split('\n')
            
            relevant_lines = []
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in keywords):
                    relevant_lines.append(line)
            
            if relevant_lines:
                content = '\n'.join(relevant_lines[:10])  # 최대 10줄
            else:
                content = f"⚠️ '{structure}' 구조의 결과를 찾을 수 없습니다."
        
        parsed_results[structure] = content
    
    return parsed_results

def create_analysis_workflow(purpose_enum, objective_enums):
    """워크플로우 생성 함수"""
    system = AnalysisSystem()
    return system.suggest_analysis_steps(purpose_enum, objective_enums)

def validate_user_inputs(user_inputs):
    """사용자 입력 검증 함수"""
    missing_fields = [field for field in REQUIRED_FIELDS if not user_inputs.get(field)]
    return missing_fields

def create_pdf_summary_dict(user_inputs, pdf_summary):
    """PDF 요약 딕셔너리 생성 함수"""
    return {
        "pdf_summary": pdf_summary,
        "project_name": user_inputs.get("project_name", ""),
        "owner": user_inputs.get("owner", ""),
        "site_location": user_inputs.get("site_location", ""),
        "site_area": user_inputs.get("site_area", ""),
        "building_type": user_inputs.get("building_type", ""),
        "project_goal": user_inputs.get("project_goal", "")
    }

def validate_and_fix_prompt(dsl_block: dict) -> dict:
    """DSL 블록 검증 및 수정"""
    content_dsl = dsl_block.get("content_dsl", {})
    output_structure = content_dsl.get("output_structure", [])
    
    if not output_structure:
        st.warning("⚠️ output_structure가 정의되지 않았습니다.")
        return dsl_block
    
    # 중복 제거
    unique_structure = list(dict.fromkeys(output_structure))
    if len(unique_structure) != len(output_structure):
        st.info("ℹ️ 중복된 구조명이 제거되었습니다.")
        content_dsl["output_structure"] = unique_structure
    
    return dsl_block

def generate_optimization_analysis(user_inputs, cot_history):
    """최적화 조건 분석 생성 함수"""
    from agent_executor import execute_agent
    
    # 분석 결과 요약
    analysis_summary = "\n\n".join([
        f"**{h['step']}**: {h.get('result', '')}"
        for h in cot_history
    ])
    
    optimization_prompt = f"""
프로젝트 정보:
- 프로젝트명: {user_inputs.get('project_name', '')}
- 건물 유형: {user_inputs.get('building_type', '')}
- 대지 위치: {user_inputs.get('site_location', '')}
- 건축주: {user_inputs.get('owner', '')}
- 대지 면적: {user_inputs.get('site_area', '')}
- 프로젝트 목표: {user_inputs.get('project_goal', '')}

분석 결과:
{analysis_summary}

위 정보를 바탕으로 매스별 최적화 조건을 자동으로 분석해주세요.

분석 요청사항:
1. **매스별 중요 프로그램 식별**: 각 매스에서 어떤 프로그램이 가장 중요한지 분석
2. **매스별 최적화 조건**: 각 매스의 특성에 따른 최적화 조건 제시
3. **프로그램별 우선순위**: 매스 내에서 프로그램들의 우선순위 분석

각 매스별로 다음 항목들을 분석해주세요:

1. **인지성**: 외부 인식, 동선 유도
2. **프라이버시**: 외부 시야 차단 필요성
3. **프로그램 연계 필요성**: 다른 공간과의 인접 배치 필요성
4. **보안성**: 출입구·코어·방문자 제어
5. **조망/채광 민감도**: 전망 확보, 자연광 필요 여부
6. **향후 확장 가능성**: 평면 또는 프로그램 확장 가능성
7. **동선 분리성**: 방문객 vs 연수생 vs 운영자
8. **구조적 유연성**: 스팬, 기둥 배치, 무주공간 등 구조 제약 수준
9. **이용 시간대 특성**: 주간/야간 사용 구분, 겹침 여부
10. **대지 조건 연계성**: 경사, 조망, 레벨차 등 대지와의 물리적 적합성

각 항목에 대해 다음을 포함해주세요:
- 목적 (purpose)
- 중요도 (importance: 높음/중간/낮음)
- 고려사항 (considerations)
- 해당 매스에서의 특별한 고려사항

분석 결과는 다음 형식으로 제시해주세요:

## 매스별 최적화 조건 분석 결과

### [매스명 1]
#### 1. 인지성
- **목적**: [목적 설명]
- **중요도**: [높음/중간/낮음]
- **고려사항**: [고려사항]
- **특별 고려사항**: [해당 매스 특성]

#### 2. 프라이버시
- **목적**: [목적 설명]
- **중요도**: [높음/중간/낮음]
- **고려사항**: [고려사항]
- **특별 고려사항**: [해당 매스 특성]

[이하 3-10번 항목 동일한 형식으로 계속...]

### [매스명 2]
[동일한 형식으로 계속...]

## 종합 최적화 전략
[전체 매스에 대한 종합적인 최적화 전략 제시]
"""
    
    return execute_agent(optimization_prompt)

def generate_narrative(user_inputs, cot_history, user_settings):
    """Narrative 생성 함수"""
    from agent_executor import execute_agent
    
    # Implementation for narrative generation
    narrative_prompt = f"""
프로젝트 정보를 바탕으로 건축설계 발표용 Narrative를 생성해주세요.
"""
    
    return execute_agent(narrative_prompt)

def generate_midjourney_prompt(user_inputs, cot_history, image_settings):
    """Midjourney 프롬프트 생성 함수"""
    from agent_executor import execute_midjourney_prompt  # ← 새로운 함수로 변경
    
    # 분석 결과 요약
    analysis_results = []
    if cot_history:
        for i, history in enumerate(cot_history, 1):
            analysis_results.append({
                'step': history.get('step', f'단계 {i}'),
                'summary': history.get('summary', ''),
                'insight': history.get('insight', ''),
                'result': history.get('result', '')
            })
    
    # 분석 결과 텍스트 생성
    if analysis_results:
        internal_analysis = "\n\n".join([
            f"**{h['step']}**: {h.get('summary', '')}"
            for h in analysis_results
        ])
    else:
        internal_analysis = ""
    
    # 외부 문서 내용과 내부 분석을 결합
    external_content = image_settings.get('external_analysis_content', '')
    if external_content and internal_analysis:
        analysis_summary = f"{external_content}\n\n**내부 분석 결과:**\n{internal_analysis}"
    elif external_content:
        analysis_summary = external_content
    elif internal_analysis:
        analysis_summary = internal_analysis
    else:
        analysis_summary = "분석 결과가 없습니다. 프로젝트 정보만을 기반으로 이미지 프롬프트를 생성합니다."
    
    # 개선된 이미지 생성 프롬프트
    image_prompt = f"""
당신은 건축 이미지 생성 전문가입니다. 분석 결과를 바탕으로 Midjourney에서 사용할 수 있는 구체적이고 효과적인 프롬프트를 생성해주세요.

##  프로젝트 정보
- 프로젝트명: {user_inputs.get('project_name', '')}
- 건물 유형: {user_inputs.get('building_type', '')}
- 대지 위치: {user_inputs.get('site_location', '')}
- 건축주: {user_inputs.get('owner', '')}
- 대지 면적: {user_inputs.get('site_area', '')}

## 분석 결과
{analysis_summary}

##  이미지 생성 요청
- 이미지 유형: {image_settings.get('image_type', '')}
- 스타일: {', '.join(image_settings.get('style_preference', [])) if image_settings.get('style_preference') else '기본'}
- 추가 설명: {image_settings.get('additional_description', '')}

##  출력 형식

**한글 설명:**
[이미지에 대한 한글 설명 - 건축적 특징, 분위기, 핵심 요소 등]

**English Midjourney Prompt:**
[구체적이고 실행 가능한 영어 프롬프트]

##  프롬프트 생성 가이드라인

**이미지 유형별 키워드:**
- **외관 렌더링**: building facade, exterior view, architectural elevation, material texture
- **내부 공간**: interior space, indoor lighting, furniture arrangement, spatial atmosphere
- **마스터플랜**: master plan, site layout, landscape design, circulation plan
- **상세도**: architectural detail, construction detail, material junction
- **컨셉 이미지**: concept visualization, mood board, artistic expression
- **조감도**: aerial view, bird's eye view, overall building form, site context

**스타일별 키워드:**
- **현대적**: modern, contemporary, clean lines, minimalist
- **미니멀**: minimal, simple, uncluttered, essential elements
- **자연친화적**: sustainable, green building, organic, eco-friendly
- **고급스러운**: luxury, premium, sophisticated, elegant
- **기능적**: functional, practical, efficient, user-friendly
- **예술적**: artistic, creative, expressive, innovative
- **상업적**: commercial, business-oriented, professional

**기술적 키워드:**
- architectural photography, professional rendering, hyperrealistic, 8k, high quality
- wide angle, natural lighting, golden hour, dramatic shadows, ambient lighting
- architectural visualization, photorealistic, modern design

**프롬프트 구조:**
[이미지 종류] + [건축 스타일] + [공간 유형] + [재료/텍스처] + [조명/분위기] + [환경/맥락] + [기술적 키워드] + [이미지 비율]

## ⚠️ 중요 지시사항
1. **분석 결과 반영**: 반드시 분석 결과의 건축적 특징을 프롬프트에 반영
2. **구체성**: 추상적이 아닌 구체적이고 실행 가능한 프롬프트 생성
3. **건축적 정확성**: 실제 건축물의 구조와 형태를 정확히 반영
4. **시각적 임팩트**: 조형적 아름다움과 상징성을 강조
5. **환경적 맥락**: 주변 환경과의 조화로운 관계 표현

위 가이드라인에 따라 한글 설명과 영어 Midjourney 프롬프트를 생성해주세요.
"""
    
    return execute_midjourney_prompt(image_prompt)  # ← 새로운 함수 호출

def render_analysis_workflow():
    """분석 워크플로우 렌더링"""
    st.header("분석 워크플로우")
    
    # 분석이 시작되었는지 확인
    if st.session_state.get('analysis_started', False):
        # 분석 실행 UI 호출
        render_analysis_execution()
        return
    
    # 사용자 입력 가져오기
    user_inputs = get_user_inputs()
    
    # 1단계: 목적과 용도 선택
    st.subheader("1단계: 분석 목적과 용도 선택")
    
    from analysis_system import AnalysisSystem, PurposeType, ObjectiveType
    system = AnalysisSystem()
    
    # 용도 선택
    purpose_options = [purpose.value for purpose in PurposeType]
    selected_purpose = st.selectbox(
        "건물 용도 선택",
        purpose_options,
        key="selected_purpose_workflow"
    )
    
    # 선택된 용도에 따른 목적 옵션 표시
    if selected_purpose:
        purpose_enum = PurposeType(selected_purpose)
        available_objectives = system.get_available_objectives(purpose_enum)
        objective_options = [obj.value for obj in available_objectives]
        
        selected_objectives = st.multiselect(
            "분석 목적 선택 (복수 선택 가능)",
            objective_options,
            key="selected_objectives_workflow"
        )
        
        # 선택된 목적들을 ObjectiveType으로 변환
        objective_enums = [ObjectiveType(obj) for obj in selected_objectives]
        
        # 워크플로우 제안
        if selected_objectives:
            st.success(f"✅ 선택된 용도: {selected_purpose}")
            st.success(f"✅ 선택된 목적: {', '.join(selected_objectives)}")
            
            # 워크플로우 생성
            workflow = system.suggest_analysis_steps(purpose_enum, objective_enums)
            
            # 사이드바에 추가 가능한 단계들 표시 (매번 새로 계산)
            st.sidebar.markdown("### ➕ 추가 선택 가능한 단계")
            
            # 프롬프트 블록 로드
            from dsl_to_prompt import load_prompt_blocks
            blocks = load_prompt_blocks()
            extra_blocks = blocks["extra"]
            
            # 현재 선택된 단계들 (실시간 계산)
            current_step_ids = set()
            if st.session_state.get('editable_steps'):
                for step in st.session_state.editable_steps:
                    current_step_ids.add(step.id)
            
            # 자동 제안된 단계들 (editable_steps 기준으로 계산)
            auto_suggested_ids = set()
            if st.session_state.get('editable_steps'):
                # editable_steps를 기준으로 auto_suggested_ids 계산
                auto_suggested_ids.update({step.id for step in st.session_state.editable_steps})
            elif workflow:
                # editable_steps가 없을 때만 workflow.steps 사용
                auto_suggested_ids.update({step.id for step in workflow.steps})
            
            # 추가 가능한 단계들 필터링
            additional_blocks = []
            for block in extra_blocks:
                block_id = block["id"]
                if "hyderabad" in block_id.lower():
                    continue
                # editable_steps에 없는 블록들만 추가 가능
                if block_id not in current_step_ids:
                    additional_blocks.append(block)
            
            st.sidebar.write(f"**디버깅**: additional_blocks 개수 = {len(additional_blocks)}")
            
            if additional_blocks:
                st.sidebar.write("**추가로 선택 가능한 단계**:")
                
                for block in additional_blocks:
                    block_id = block["id"]
                    
                    # 선택 가능한 단계
                    if st.sidebar.button(f"➕ {block['title']}", key=f"add_block_{block_id}_workflow"):
                        # 단계 추가
                        from analysis_system import AnalysisStep
                        
                        # 권장 순서에 따른 적절한 위치 찾기
                        cot_order = system._load_recommended_cot_order()
                        new_step_order = cot_order.get(block_id, 999)  # 기본값을 높게 설정
                        
                        new_step = AnalysisStep(
                            id=block_id,
                            title=block['title'],
                            description=block.get('description', ''),
                            is_optional=True,
                            order=new_step_order,
                            category="추가 단계"
                        )
                        
                        # editable_steps에 추가
                        if 'editable_steps' not in st.session_state:
                            st.session_state.editable_steps = []
                        
                        st.session_state.editable_steps.append(new_step)
                        
                        # 권장 순서로 재정렬
                        sorted_steps = system.sort_steps_by_recommended_order(st.session_state.editable_steps)
                        for i, step in enumerate(sorted_steps, 1):
                            step.order = i
                        
                        st.session_state.editable_steps = sorted_steps
                        
                        # 성공 메시지 표시
                        st.sidebar.success(f"'{block['title']}' 단계가 추가되었습니다!")
                        st.rerun()
            else:
                st.sidebar.info("모든 관련 단계가 자동으로 선택되었습니다.")
            
            # 제안된 단계들 표시 및 편집 기능
            st.subheader("2단계: 분석 단계 편집")
            st.info("제안된 단계들을 자유롭게 편집할 수 있습니다:")
            
            # 편집 가능한 단계 리스트 업데이트 (목적 변경 시마다)
            current_selection = f"{selected_purpose}_{','.join(selected_objectives)}"
            if 'current_selection' not in st.session_state or st.session_state.current_selection != current_selection:
                st.session_state.editable_steps = workflow.steps.copy()
                st.session_state.current_selection = current_selection
            
            # 단계 편집 인터페이스
            st.markdown("#### 현재 분석 단계")
            
            # 각 단계를 편집 가능한 형태로 표시
            for i, step in enumerate(st.session_state.editable_steps):
                with st.expander(f"{i+1}. {step.title}", expanded=True):
                    col_a, col_b, col_c = st.columns([2, 1, 1])
                    
                    with col_a:
                        st.markdown(f"**설명**: {step.description}")
                        if step.is_required:
                            st.caption("🔴 필수 단계")
                        elif step.is_recommended:
                            st.caption("🟡 권장 단계")
                        else:
                            st.caption("🟢 선택 단계")
                    
                    with col_b:
                        if st.button("❌ 제거", key=f"remove_{step.id}_workflow"):
                            st.session_state.editable_steps.pop(i)
                            st.session_state.sidebar_updated = True
                            st.rerun()
                    
                    with col_c:
                        if i > 0:
                            if st.button("⬆️ 위로", key=f"up_{step.id}_workflow"):
                                st.session_state.editable_steps[i], st.session_state.editable_steps[i-1] = \
                                    st.session_state.editable_steps[i-1], st.session_state.editable_steps[i]
                                st.rerun()
                        if i < len(st.session_state.editable_steps) - 1:
                            if st.button("⬇️ 아래로", key=f"down_{step.id}_workflow"):
                                st.session_state.editable_steps[i], st.session_state.editable_steps[i+1] = \
                                    st.session_state.editable_steps[i+1], st.session_state.editable_steps[i]
                                st.rerun()
            
            # 워크플로우 편집 완료 후 제어 버튼들 표시 (통합)
            st.markdown("---")
            
            # 제어 버튼들 (타이틀 없이)
            col1, col2 = st.columns([1, 1])
            
            with col1:
                if st.button("🔄 권장 순서 제안", type="secondary", help="선택된 단계들을 권장 CoT 순서로 재정렬합니다", key="recommend_order_workflow"):
                    from analysis_system import AnalysisSystem
                    system = AnalysisSystem()
                    
                    # 현재 단계들을 권장 순서로 정렬 (editable_steps 사용)
                    sorted_steps = system.sort_steps_by_recommended_order(st.session_state.editable_steps)
                    
                    # 순서 번호 업데이트
                    for i, step in enumerate(sorted_steps, 1):
                        step.order = i
                    
                    # editable_steps를 직접 업데이트 (메인 편집 인터페이스에 반영)
                    st.session_state.editable_steps = sorted_steps
                    st.success("✅ 단계가 권장 순서로 재정렬되었습니다!")
                    st.rerun()
            
            with col2:
                if st.button("분석 시작", type="primary", help="선택된 단계들로 분석을 시작합니다", key="start_analysis_workflow"):
                    # 분석 시작 시 editable_steps를 workflow_steps로 복사
                    st.session_state.workflow_steps = st.session_state.editable_steps.copy()
                    st.session_state.analysis_started = True
                    # current_step_index를 0으로 초기화하지 않고 기존 값 유지
                    if 'current_step_index' not in st.session_state:
                        st.session_state.current_step_index = 0
                    st.success("분석이 시작되었습니다!")
                    st.rerun()

def generate_optimization_analysis_with_external_content(user_inputs, cot_history, analysis_summary):
    """외부 문서 내용을 포함한 최적화 조건 분석 생성 함수 - 개선된 버전"""
    from agent_executor import execute_agent
    
    optimization_prompt = f"""
프로젝트 정보:
- 프로젝트명: {user_inputs.get('project_name', '')}
- 건물 유형: {user_inputs.get('building_type', '')}
- 대지 위치: {user_inputs.get('site_location', '')}
- 건축주: {user_inputs.get('owner', '')}
- 대지 면적: {user_inputs.get('site_area', '')}
- 프로젝트 목표: {user_inputs.get('project_goal', '')}

분석 결과 (내부 분석 + 외부 문서):
{analysis_summary}

위 정보를 바탕으로 각 매스별로 **가장 중요한 최적화 조건 3-5개만** 선별하여 분석해주세요.

## 분석 지침:

### 1. 매스별 특성 분석
각 매스의 **핵심 기능과 특성**을 먼저 파악하고, 그에 따라 **우선순위가 높은 항목들만** 선별하세요.

### 2. 선별 기준
- **해당 매스의 핵심 기능과 직접 관련된 항목**
- **프로젝트 성공에 결정적 영향을 미치는 항목**
- **다른 매스와 차별화되는 고유한 요구사항**
- **실제 설계 과정에서 반드시 고려해야 할 항목**

### 3. 분석 형식
각 매스별로 **3-5개의 핵심 항목만** 선택하여 다음 형식으로 분석:

## 매스별 최적화 조건 분석 결과

### [매스명]
#### 1. [선택된 항목명] - **핵심 우선순위**
- **선택 이유**: 이 매스에서 이 항목이 중요한 이유
- **목적**: 구체적인 목표
- **중요도**: 높음/중간/낮음
- **핵심 고려사항**: 2-3개의 구체적 고려사항
- **매스별 특화 전략**: 이 매스만의 고유한 접근 방법

#### 2. [선택된 항목명] - **핵심 우선순위**
[동일한 형식으로 계속...]

### 4. 선별 예시
**연구시설 매스**의 경우:
- 보안성 (R&D 기밀 보호가 핵심)
- 구조적 유연성 (연구 장비 변경 대응)
- 프라이버시 (연구 활동 보호)
- 프로그램 연계 필요성 (기존 시설과의 연결)

**업무시설 매스**의 경우:
- 조망/채광 민감도 (업무 효율성)
- 동선 분리성 (업무 집중도)
- 향후 확장 가능성 (조직 성장 대응)

**공용시설 매스**의 경우:
- 인지성 (캠퍼스 중심 역할)
- 프로그램 연계 필요성 (연결점 역할)
- 이용 시간대 특성 (다양한 사용자 대응)

## 종합 최적화 전략
각 매스의 핵심 조건들을 종합하여 **3-4개의 핵심 전략**으로 정리:

1. **[전략명]**: 구체적 실행 방안
2. **[전략명]**: 구체적 실행 방안
3. **[전략명]**: 구체적 실행 방안

위 지침에 따라 각 매스별로 **가장 중요한 3-5개 항목만** 선별하여 분석해주세요.
"""
    
    return execute_agent(optimization_prompt)

def execute_analysis_step(step_id: str, full_prompt: str) -> str:
    """분석 단계 실행"""
    # 기존 블록들
    if step_id == "document_analyzer":
        return run_document_analyzer(full_prompt)
    elif step_id == "requirement_analyzer":
        return run_requirement_analyzer(full_prompt)
    # ... existing blocks ...
    
    # 하이데라바드 프로젝트 전용 블록들 (일시적으로 비활성화)
    # elif step_id == "hyderabad_campus_expansion_analysis":
    #     return run_hyderabad_campus_expansion_analysis(full_prompt)
    # elif step_id == "hyderabad_research_infra_strategy":
    #     return run_hyderabad_research_infra_strategy(full_prompt)
    # elif step_id == "hyderabad_talent_collaboration_infra":
    #     return run_hyderabad_talent_collaboration_infra(full_prompt)
    # elif step_id == "hyderabad_welfare_branding_environment":
    #     return run_hyderabad_welfare_branding_environment(full_prompt)
    # elif step_id == "hyderabad_security_zoning_plan":
    #     return run_hyderabad_security_zoning_plan(full_prompt)
    # elif step_id == "hyderabad_masterplan_roadmap":
    #     return run_hyderabad_masterplan_roadmap(full_prompt)
    
    else:
        return f"⚠️ 알 수 없는 분석 단계: {step_id}"

def main():
    """메인 함수"""
    render_tabbed_interface()

if __name__ == "__main__":
    main() 